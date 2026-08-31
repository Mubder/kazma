"""Three dark-logo letterhead mockups in IBM Plex Sans Arabic.

Does not change production THEME. Review-only.
"""

from __future__ import annotations

import os
import shutil
import tempfile
from contextlib import contextmanager
from copy import deepcopy
from datetime import datetime
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Mm, Pt, RGBColor
from PIL import Image

from kazma_core.documents.binaries import run_soffice_cli
from kazma_core.documents.content_model import (
    BodyBlock,
    ContentModel,
    HeadingBlock,
    TableBlock,
    TitleBlock,
)
from kazma_core.documents.engines.docx import DocxEngine
from kazma_core.documents.profile import DocProfile
from kazma_core.documents.style_theme import THEME, format_document_date

HERE = Path(__file__).resolve().parent
COPY = Path(r"C:\Users\balfa\Desktop\TestingMocks.txt")
LOGO_SRC = Path(r"C:\Users\balfa\Downloads\Kazma New Logos\KazmaLogoDarkShadow.png")
FONT_SRC = HERE / "_fonts"
ASSETS = HERE / "_assets"
FACE = "IBM Plex Sans Arabic"

NAVY = "16223A"
ROYAL = "3B82F6"
SKY = "38BDF8"
INK = "16223A"
MUTED = "94A3B8"
WHITE = "FFFFFF"

PAGE_W = Mm(210)

_GREGORIAN_MONTH_AR = (
    "يناير", "فبراير", "مارس", "أبريل", "مايو", "يونيو",
    "يوليو", "أغسطس", "سبتمبر", "أكتوبر", "نوفمبر", "ديسمبر",
)
_GREGORIAN_MONTH_EN = (
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
)


def _format_stamp(when: datetime, *, rtl: bool) -> str:
    """Generation timestamp for chrome. AR: time + Gregorian + Hijri. EN: time + Gregorian."""
    clock = when.strftime("%H:%M")
    d = when.date()
    if rtl:
        greg = f"{d.day} {_GREGORIAN_MONTH_AR[d.month - 1]} {d.year} م"
        hijri = format_document_date(d, rtl=True, calendar="islamic-umalqura")
        return f"{clock}  ·  {greg}  ·  {hijri}"
    greg = f"{d.day} {_GREGORIAN_MONTH_EN[d.month - 1]} {d.year}"
    return f"{clock}  ·  {greg}"
MARGIN = Cm(1.8)
CONTENT_W = Emu(int(PAGE_W) - 2 * int(MARGIN))


def _copy_paras() -> list[str]:
    raw = COPY.read_text(encoding="utf-8")
    return [p.strip() for p in raw.replace("\r\n", "\n").split("\n") if p.strip()]


def _stage_plex() -> Path:
    for name in ("IBMPlexSansArabic-Regular.ttf", "IBMPlexSansArabic-Bold.ttf"):
        if not (FONT_SRC / name).is_file():
            raise FileNotFoundError(FONT_SRC / name)
    return FONT_SRC


def _install_user_fonts(font_dir: Path) -> None:
    dest_dir = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "Windows" / "Fonts"
    dest_dir.mkdir(parents=True, exist_ok=True)
    try:
        import ctypes
        import winreg
    except ImportError:
        return
    try:
        key = winreg.CreateKey(
            winreg.HKEY_CURRENT_USER,
            r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts",
        )
    except OSError:
        key = None
    for ttf in sorted(font_dir.glob("*.ttf")):
        dest = dest_dir / ttf.name
        if not dest.exists() or dest.stat().st_size != ttf.stat().st_size:
            shutil.copy2(ttf, dest)
        try:
            ctypes.windll.gdi32.AddFontResourceW(str(dest))
        except Exception:
            pass
        if key is not None:
            try:
                winreg.SetValueEx(key, f"{ttf.stem} (TrueType)", 0, winreg.REG_SZ, str(dest))
            except OSError:
                pass
    if key is not None:
        winreg.CloseKey(key)


@contextmanager
def _patched_theme(**kwargs):
    saved = {k: THEME[k] for k in kwargs if k in THEME}
    extras = [k for k in kwargs if k not in THEME]
    THEME.update(kwargs)
    try:
        yield
    finally:
        for k in extras:
            THEME.pop(k, None)
        THEME.update(saved)


def _prepare_logo() -> Path:
    """Crop empty canvas only. Keep the PNG alpha — never flatten onto a fill."""
    ASSETS.mkdir(exist_ok=True)
    im = Image.open(LOGO_SRC).convert("RGBA")
    alpha = im.split()[3].point(lambda v: 255 if v > 10 else 0)
    bbox = alpha.getbbox()
    if bbox is None:
        raise RuntimeError("dark logo has no visible pixels")
    pad = 16
    crop = im.crop((
        max(0, bbox[0] - pad),
        max(0, bbox[1] - pad),
        min(im.width, bbox[2] + pad),
        min(im.height, bbox[3] + pad),
    ))
    dest = ASSETS / "logo-dark.png"
    crop.save(dest, "PNG")
    return dest


def _rgb(hex6: str) -> RGBColor:
    h = hex6.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade(cell, hex6: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex6)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _vcenter(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    el = tc_pr.find(qn("w:vAlign"))
    if el is None:
        el = OxmlElement("w:vAlign")
        tc_pr.append(el)
    el.set(qn("w:val"), "center")


def _cell_margins(cell, *, twips: int = 80) -> None:
    _cell_margins_sides(
        cell,
        left=Emu(twips * 635),
        right=Emu(twips * 635),
        top=Emu(twips * 635),
        bottom=Emu(twips * 635),
    )


def _cell_margins_sides(cell, *, left, right, top, bottom) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(_twips(val)))
        node.set(qn("w:type"), "dxa")


def _no_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tc_pr.append(borders)


def _no_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")


def _tbl_width(table, width) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width.twips)))
    tbl_w.set(qn("w:type"), "dxa")


def _twips(val) -> int:
    return int(val.twips) if hasattr(val, "twips") else int(Emu(int(val)).twips)


def _tbl_cell_mar(table, *, left=0, right=0, top=0, bottom=0) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblCellMar"))
    if existing is not None:
        tbl_pr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(_twips(val)))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tbl_pr.append(mar)


def _place_in_cell(cell, table) -> None:
    """Move a body table into a cell so a bleed shell can wrap inset content."""
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    tc.append(tbl)
    tc.append(OxmlElement("w:p"))


def _bleed_shell(document, *, fill: str, height) -> object:
    """Full-page-width shaded table. Page L/R margins must already be 0."""
    table = document.add_table(rows=1, cols=1)
    _no_borders(table)
    _tbl_width(table, PAGE_W)
    _col_widths(table, [PAGE_W])
    _tbl_cell_mar(table, left=0, right=0, top=0, bottom=0)
    cell = table.cell(0, 0)
    _shade(cell, fill)
    _vcenter(cell)
    _no_cell_borders(cell)
    _cell_margins(cell, twips=0)
    _row_height(table.rows[0], height, exact=True)
    return table


def _tbl_indent(table, indent) -> None:
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(int(indent.twips)))
    ind.set(qn("w:type"), "dxa")


def _col_widths(table, widths) -> None:
    table.autofit = False
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
    else:
        grid = OxmlElement("w:tblGrid")
        table._tbl.tblPr.addnext(grid)
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.twips)))
        grid.append(gc)
    for row in table.rows:
        for cell, w in zip(row.cells, widths, strict=False):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(w.twips)))
            tc_w.set(qn("w:type"), "dxa")


def _row_height(row, height, *, exact: bool = False) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_h = tr_pr.find(qn("w:trHeight"))
    if tr_h is None:
        tr_h = OxmlElement("w:trHeight")
        tr_pr.append(tr_h)
    tr_h.set(qn("w:val"), str(int(height.twips)))
    tr_h.set(qn("w:hRule"), "exact" if exact else "atLeast")


def _style_run(run, *, size: float, color: str, bold: bool = False, rtl: bool = True) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FACE
    run.font.color.rgb = _rgb(color)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), FACE)
    if rtl:
        if r_pr.find(qn("w:rtl")) is None:
            r_pr.append(OxmlElement("w:rtl"))
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:bidi"), "ar-SA")
        lang.set(qn("w:val"), "ar-SA")
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    half = str(int(size * 2))
    sz.set(qn("w:val"), half)
    szcs = r_pr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        r_pr.append(szcs)
    szcs.set(qn("w:val"), half)


def _rtl_para(p, *, align: str = "start") -> None:
    # Under w:bidi, "start" is the physical right. Never use jc=right for Arabic.
    mapping = {
        "start": WD_ALIGN_PARAGRAPH.RIGHT,
        "end": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }
    p.alignment = mapping.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), align)


def _clear_cell(cell) -> None:
    for p in cell.paragraphs:
        p.text = ""


def _put_logo(cell, path: Path, width: Emu) -> None:
    p = cell.paragraphs[0]
    p.text = ""
    _rtl_para(p, align="center")
    run = p.add_run()
    run.add_picture(str(path), width=width)


def _chrome_para(p, *, align: str, rtl: bool) -> None:
    if rtl:
        _rtl_para(p, align=align)
        return
    mapping = {
        "start": WD_ALIGN_PARAGRAPH.LEFT,
        "end": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }
    p.alignment = mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "0")
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc_val = "center" if align == "center" else ("start" if align in ("start", "left") else "end")
    jc.set(qn("w:val"), jc_val)


def _put_text(
    cell,
    lines: list[tuple[str, float, str, bool]],
    *,
    align: str = "start",
    rtl: bool = True,
) -> None:
    first = True
    for text, size, color, bold in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.text = ""
        _chrome_para(p, align=align, rtl=rtl)
        run = p.add_run(text)
        _style_run(run, size=size, color=color, bold=bold, rtl=rtl)


def _prepend_element(document, element) -> None:
    body = document.element.body
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    body.insert(0, element)


def _insert_after(anchor, element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    anchor.addnext(element)


def _bar_table(document, *, fill: str, height, width, indent) -> object:
    table = document.add_table(rows=1, cols=1)
    _no_borders(table)
    _tbl_width(table, width)
    _tbl_indent(table, indent)
    _col_widths(table, [width])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    _no_cell_borders(cell)
    _clear_cell(cell)
    _cell_margins(cell, twips=0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")
    run = p.add_run(" ")
    run.font.size = Pt(1)
    _row_height(table.rows[0], height, exact=True)
    return table


def _tbl_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _full_row(document, widths, *, height, fill: str | None) -> object:
    """One fixed-layout row the full page width (page L/R must be 0)."""
    table = document.add_table(rows=1, cols=len(widths))
    _no_borders(table)
    _tbl_width(table, PAGE_W)
    _col_widths(table, widths)
    _tbl_fixed(table)
    _tbl_cell_mar(table, left=0, right=0, top=0, bottom=0)
    if fill:
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is not None:
            for el in borders:
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), fill)
    _row_height(table.rows[0], height)
    for cell in table.rows[0].cells:
        _vcenter(cell)
        _cell_margins(cell, twips=40)
        _no_cell_borders(cell)
        if fill:
            _shade(cell, fill)
    return table


def _letterhead_a(document, logo: Path) -> None:
    """Full-bleed navy band; logo + type inset to the body 1.8 cm margin.

    Grid is [pad | logo | brand | meta | pad]. LibreOffice mirrors RTL tables,
    so the visual order is pad / date / brand / logo / pad — logo on the right,
    both pads 1.8 cm to match the body.
    """
    logo_w = Cm(1.55)
    col_logo = Cm(2.0)
    col_meta = Cm(3.6)
    col_brand = Emu(int(CONTENT_W) - int(col_logo) - int(col_meta))
    widths = [MARGIN, col_logo, col_brand, col_meta, MARGIN]
    table = _full_row(document, widths, height=Cm(2.3), fill=NAVY)
    _put_logo(table.cell(0, 1), logo, logo_w)
    _put_text(
        table.cell(0, 2),
        [
            ("كاظمه", 16, WHITE, True),
            ("منظومة كاظمة للذكاء الاصطناعي", 9, SKY, False),
        ],
        align="start",
    )
    _put_text(
        table.cell(0, 3),
        [
            ("٣١ أغسطس ٢٠٢٦", 8.5, MUTED, False),
            ("رسالة رسمية", 8.5, SKY, False),
        ],
        align="end",
    )
    stripe = _bar_table(document, fill=ROYAL, height=Pt(3.5), width=PAGE_W, indent=Emu(0))
    _tbl_fixed(stripe)
    _prepend_element(document, stripe._tbl)
    _prepend_element(document, table._tbl)


def _letterhead_b(document, logo: Path, *, rtl: bool, chrome: dict, stamp: str) -> None:
    """White paper, transparent K, navy wordmark, royal + sky rules."""
    short = str(chrome.get("brand_short") or ("كاظمه" if rtl else "Kazma"))
    long = str(chrome.get("brand") or ("منظومة كاظمة للذكاء الاصطناعي" if rtl else "Kazma AI Platform"))
    logo_w = Cm(1.7)
    col_logo = Cm(2.1)
    col_brand = Emu(int(CONTENT_W) - int(col_logo))
    table = document.add_table(rows=1, cols=2)
    _no_borders(table)
    _tbl_width(table, CONTENT_W)
    _tbl_fixed(table)
    _tbl_cell_mar(table, left=0, right=0, top=0, bottom=0)
    _row_height(table.rows[0], Cm(2.35))
    for cell in table.rows[0].cells:
        _vcenter(cell)
        _cell_margins(cell, twips=40)
        _no_cell_borders(cell)
    lines = [(short, 18, INK, True), (long, 10, MUTED, False), (stamp, 8, MUTED, False)]
    if rtl:
        # LO mirrors RTL tables → wordmark left, logo right.
        _col_widths(table, [col_logo, col_brand])
        _put_logo(table.cell(0, 0), logo, logo_w)
        _put_text(table.cell(0, 1), lines, align="end", rtl=True)
    else:
        _col_widths(table, [col_brand, col_logo])
        _put_text(table.cell(0, 0), lines, align="start", rtl=False)
        _put_logo(table.cell(0, 1), logo, logo_w)
    royal = _bar_table(document, fill=ROYAL, height=Pt(3.2), width=CONTENT_W, indent=Emu(0))
    sky = _bar_table(document, fill=SKY, height=Pt(1.6), width=CONTENT_W, indent=Emu(0))
    _prepend_element(document, sky._tbl)
    _prepend_element(document, royal._tbl)
    _prepend_element(document, table._tbl)


def _letterhead_c(document, logo: Path) -> None:
    """Tall navy masthead; larger transparent K; type inset to 1.8 cm.

    Grid [pad | logo | brand | pad] → LO mirror → pad / brand / logo / pad.
    """
    logo_w = Cm(2.35)
    col_logo = Cm(2.8)
    col_brand = Emu(int(CONTENT_W) - int(col_logo))
    widths = [MARGIN, col_logo, col_brand, MARGIN]
    table = _full_row(document, widths, height=Cm(3.4), fill=NAVY)
    _put_logo(table.cell(0, 1), logo, logo_w)
    _put_text(
        table.cell(0, 2),
        [
            ("كاظمه", 22, WHITE, True),
            ("منظومة كاظمة للذكاء الاصطناعي", 11, SKY, False),
            ("Kazma AI Platform", 9, MUTED, False),
        ],
        align="end",
    )
    stripe = _bar_table(document, fill=SKY, height=Pt(4.2), width=PAGE_W, indent=Emu(0))
    _tbl_fixed(stripe)
    _prepend_element(document, stripe._tbl)
    _prepend_element(document, table._tbl)


def _spacer_after(document, anchor_tbl, *, pt: float = 14) -> None:
    p = document.add_paragraph("")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)
    _insert_after(anchor_tbl, p._p)


def _zero_header_para(section) -> None:
    """An empty header paragraph still eats a Normal-style line (~16 pt)."""
    for p in section.header.paragraphs:
        p.text = ""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p_pr = p._p.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "0")
        spacing.set(qn("w:lineRule"), "exact")
    sect_pr = section._sectPr
    for tag in ("w:headerReference",):
        for el in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(el)


def _set_para_ind(p_el, left, right) -> None:
    p_pr = p_el.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_el.insert(0, p_pr)
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(_twips(left)))
    ind.set(qn("w:right"), str(_twips(right)))


def _inset_body(document) -> None:
    """When page L/R are 0 for a bleed band, restore 1.8 cm on body + footer."""
    body = document.element.body
    past_chrome = False
    for child in body:
        if not past_chrome and child.tag == qn("w:tbl"):
            continue
        past_chrome = True
        if child.tag == qn("w:p"):
            _set_para_ind(child, MARGIN, MARGIN)
    for section in document.sections:
        for p in section.footer.paragraphs:
            _set_para_ind(p._p, MARGIN, MARGIN)


def _set_page(document, *, bleed: bool) -> None:
    for section in document.sections:
        section.header_distance = Cm(0)
        _zero_header_para(section)
        if bleed:
            section.top_margin = Cm(0)
            section.left_margin = Cm(0)
            section.right_margin = Cm(0)
        else:
            section.top_margin = MARGIN
            section.left_margin = MARGIN
            section.right_margin = MARGIN
    if bleed:
        _inset_body(document)


def _model(paras: list[str], *, footer: str) -> ContentModel:
    blocks: list = [TitleBlock(text="أبرز ميزات كاظمه", level=0)]
    if paras:
        blocks.append(BodyBlock(text=paras[0]))
    if len(paras) > 1:
        blocks.append(HeadingBlock(text="كاظمه تتحدث MCP", level=1))
        blocks.append(BodyBlock(text=paras[1]))
    return ContentModel(
        blocks=blocks,
        header="",
        footer=footer,
        page_numbers=True,
        author="Kazma",
        subject="Letterhead mockup — IBM Plex Sans Arabic",
    )


def _raster(pdf: Path) -> None:
    doc = fitz.open(pdf)
    pix = doc[0].get_pixmap(matrix=fitz.Matrix(1.7, 1.7), alpha=False)
    pix.save(str(pdf.with_suffix(".png")))
    doc.close()


def _pdf_fonts(pdf: Path) -> list[str]:
    doc = fitz.open(pdf)
    names = sorted({row[3] for row in doc[0].get_fonts()})
    doc.close()
    return names


def _docx_to_pdf(docx: Path, pdf: Path, font_dir: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="kazma_lh_") as tmp:
        tmp_dir = Path(tmp)
        fonts_dst = tmp_dir / "user" / "fonts"
        fonts_dst.mkdir(parents=True)
        for ttf in font_dir.glob("*.ttf"):
            shutil.copy2(ttf, fonts_dst / ttf.name)
        work = tmp_dir / docx.name
        shutil.copy2(docx, work)
        proc = run_soffice_cli(
            (
                "--headless",
                "--nologo",
                "--norestore",
                f"-env:UserInstallation={tmp_dir.as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_dir),
                str(work),
            ),
            timeout=120,
            cwd=tmp_dir,
        )
        produced = tmp_dir / (docx.stem + ".pdf")
        if not produced.is_file():
            raise RuntimeError(
                f"LibreOffice did not produce PDF (code={proc.returncode})\n"
                f"stdout={proc.stdout}\nstderr={proc.stderr}"
            )
        shutil.move(str(produced), str(pdf))


def _write_variant(
    *,
    stem: str,
    paras: list[str],
    footer: str,
    logo: Path,
    builder,
    bleed: bool,
    font_dir: Path,
) -> None:
    blob = "\n\n".join(paras)
    profile = DocProfile.for_content(blob, language="ar", rtl=True)
    profile.theme = dict(THEME)
    profile.theme["font_arabic"] = FACE
    profile.theme["font_latin"] = FACE
    profile.theme["body_size_ar"] = 12.0
    profile.theme["line_height_ar"] = 1.65

    docx = HERE / f"{stem}.docx"
    pdf = HERE / f"{stem}.pdf"
    DocxEngine(profile).render(_model(paras, footer=footer), docx)

    document = Document(str(docx))
    builder(document, logo)
    _set_page(document, bleed=bleed)
    last_chrome = None
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            last_chrome = child
            continue
        break
    if last_chrome is not None:
        _spacer_after(document, last_chrome, pt=12)
    document.save(str(docx))

    _docx_to_pdf(docx, pdf, font_dir)
    _raster(pdf)
    print(stem, "pdf-fonts=", _pdf_fonts(pdf))


def _model_long() -> ContentModel:
    """Three-page bilingual sample on stationery B. Real product copy, not lorem."""
    paras = _copy_paras()
    ar_open = paras[0] if paras else ""
    ar_mcp = paras[1] if len(paras) > 1 else ""
    blocks: list = [
        TitleBlock(text="أبرز ميزات كاظمه", level=0),
        BodyBlock(text=ar_open),
        BodyBlock(text=(
            "Kazma is a multi-platform AI agent: one supervisor brain, swarm "
            "workers, and the same tools on Telegram, Discord, Slack, the web "
            "app, and the TUI. This file is a three-page proof of stationery B "
            "— IBM Plex Sans Arabic for Arabic, the same family for English. "
            "Page 1 carries the full letterhead. Pages 2 and 3 keep a small "
            "running header and the footer so a long letter still looks like "
            "the same document."
        )),
        HeadingBlock(text="كاظمه تتحدث MCP", level=1),
        BodyBlock(text=ar_mcp),
        BodyBlock(text=(
            "MCP is the open standard that connects a model to tools. Point "
            "Kazma at any MCP server — databases, design tools, internal APIs, "
            "custom automation — and the agent calls those tools under the same "
            "HITL and commitment gates as native skills. Latin tokens such as "
            "MCP and API stay left-to-right inside the Arabic line."
        )),
        HeadingBlock(text="الذاكرة المعرفية", level=1),
        BodyBlock(text=(
            "الذاكرة في كاظمه ليست نافذة محادثة تُنسى عند إغلاق التبويب، "
            "بل محرك معرفي يحتفظ بالمعتقدات والوقائع والالتزامات عبر الجلسات، "
            "ويعتمد مصدراً موثوقاً فلا يُسمح لنصٍ استنتجه النموذج أن يكتب فوق "
            "ما أثبته المستخدم. النسخ الاحتياطي والتصدير يعملان على إيقاع ثابت، "
            "حيث قاعدة الذاكرة الساخنة معزولة عن طابور الصيانة حتى لا تتوقف "
            "القراءة أثناء الكتابة في الخلفية، وهذا هو الفرق بين وكيلٍ يتذكر "
            "فعلاً وآخر يعيد طرح الأسئلة في كل صباح."
        )),
        BodyBlock(text=(
            "Memory is not the last twenty turns of a chat. Kazma stores "
            "beliefs, episodes, and commitments in a durable engine: hot reads "
            "on one database, maintenance on another, so a backup job cannot "
            "stall recall. A user-asserted fact is not overwritten by a guess "
            "the model invented mid-sentence. That is the class of failure this "
            "layer exists to stop — the invented date that got scheduled over "
            "the real one."
        )),
        HeadingBlock(text="لمحة", level=1),
        TableBlock(
            headers=["Capability", "القدرة", "Why it matters"],
            rows=[
                [
                    "Cognitive memory",
                    "ذاكرة معرفية",
                    "Remembers across sessions; user facts win over model guesses.",
                ],
                [
                    "MCP",
                    "خوادم الأدوات",
                    "Plug in external tools without a parallel unsafe path.",
                ],
                [
                    "HITL gates",
                    "بوابات الموافقة",
                    "Shell, writes, and outbound posts pause until you approve.",
                ],
                [
                    "Swarm",
                    "السرب",
                    "Named workers, handoff limits, and a durable task store.",
                ],
            ],
        ),
        HeadingBlock(text="التشغيل اليومي", level=1),
        BodyBlock(text=(
            "التذكيرات تُجدول مع هدف التسليم وقت الإنشاء لا وقت الإطلاق، حتى "
            "لا تضيع بعد خمس دقائق من انتهاء الجلسة، وإشعارات دورة الحياة تصل "
            "إلى المنصات عند الإقلاع والإيقاف والفشل، فيعرف المشغّل من المحادثة "
            "أن الخادم عاد أو أنه عالق في الإقلاع. النسخ الاحتياطي الليلي "
            "يشمل SQLite وPostgres والجداول التي تملكها كاظمه فقط، حتى لا "
            "يُستعاد شيء من تطبيق آخر على القاعدة المشتركة."
        )),
        BodyBlock(text=(
            "Reminders capture the delivery target when you schedule them, not "
            "when they fire. Lifecycle notices land on every configured chat "
            "when the server starts, restarts, or fails to boot. Nightly backup "
            "is table-filtered on Postgres so a shared database neither leaks a "
            "neighbour's data in nor restores over it."
        )),
        HeadingBlock(text="ختام", level=1),
        BodyBlock(text=(
            "هذه رسالة تجريبية من ثلاث صفحات لاختبار القرطاسية ب، والنص عربي "
            "وإنجليزي بالتناوب حتى يظهر اتجاه كل كتلة. وكيلك. بياناتك. جهازك."
        )),
        BodyBlock(text=(
            "A three-page stationery proof, not a product spec. Page 1 is the "
            "full letterhead; pages 2–3 keep the running header. "
            "Your agent. Your data. Your machine."
        )),
    ]
    return ContentModel(
        blocks=blocks,
        header="",
        footer=None,
        page_numbers=True,
        author="Kazma",
        subject="Letterhead B — 3-page EN/AR proof",
    )


def _copy_element_children(src, dst) -> None:
    src_el = src._element
    dst_el = dst._element
    for child in list(dst_el):
        dst_el.remove(child)
    for child in src_el:
        dst_el.append(deepcopy(child))


def _empty_header(header) -> None:
    for p in header.paragraphs:
        p.text = ""
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p_pr = p._p.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "20")
        spacing.set(qn("w:lineRule"), "exact")


def _add_running_header(header, *, rtl: bool, chrome: dict, stamp: str) -> None:
    """Continuation chrome: brand, then timestamp, then a royal rule."""
    short = str(chrome.get("brand_short") or ("كاظمه" if rtl else "Kazma"))
    brand_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    brand_p.text = ""
    _chrome_para(brand_p, align="start", rtl=rtl)
    brand_run = brand_p.add_run(short)
    _style_run(brand_run, size=10, color=INK, bold=True, rtl=rtl)

    stamp_p = header.add_paragraph()
    # Page 1 puts the stamp under the brand on the physical left (logo is right).
    # EN start = left; AR end = physical left.
    _chrome_para(stamp_p, align="start" if not rtl else "end", rtl=rtl)
    stamp_run = stamp_p.add_run(stamp)
    _style_run(stamp_run, size=8, color=MUTED, bold=False, rtl=rtl)
    p_pr = stamp_p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), ROYAL)


def _raster_all(pdf: Path) -> None:
    doc = fitz.open(pdf)
    n = doc.page_count
    for i in range(n):
        pix = doc[i].get_pixmap(matrix=fitz.Matrix(1.55, 1.55), alpha=False)
        pix.save(str(pdf.with_name(f"{pdf.stem}-p{i + 1}.png")))
    doc.close()
    print(pdf.stem, "pages=", n)


def _model_long_en() -> ContentModel:
    """English stationery B proof — chrome must be English, page label 'Page N'."""
    blocks: list = [
        TitleBlock(text="Kazma product note", level=0),
        BodyBlock(text=(
            "Kazma is a multi-platform AI agent: one supervisor brain, swarm "
            "workers, and the same tools on Telegram, Discord, Slack, the web "
            "app, and the TUI. This file is the English twin of stationery B. "
            "Header, footer, and page label must read in English because the "
            "document language is English — not because a second template exists."
        )),
        HeadingBlock(text="Memory and tools", level=1),
        BodyBlock(text=(
            "Memory is not the last twenty turns of a chat. Kazma stores "
            "beliefs, episodes, and commitments in a durable engine: hot reads "
            "on one database, maintenance on another, so a backup job cannot "
            "stall recall. A user-asserted fact is not overwritten by a guess "
            "the model invented mid-sentence."
        )),
        BodyBlock(text=(
            "MCP is the open standard that connects a model to tools. Point "
            "Kazma at any MCP server — databases, design tools, internal APIs, "
            "custom automation — and the agent calls those tools under the same "
            "HITL gates as native skills."
        )),
        HeadingBlock(text="What to check", level=1),
        BodyBlock(text=(
            "Page 1 carries the full letterhead with the short brand 'Kazma' "
            "and the long footer 'Kazma AI Platform'. Page 2 keeps a running "
            "header and a footer that says Page 2 — not a bare digit, and not "
            "the Arabic chrome. If you still see منظومة or صفحة, language "
            "routing failed."
        )),
        BodyBlock(text=(
            "This second page exists so the continuation header and the live "
            "PAGE field can be judged. Your agent. Your data. Your machine."
        )),
    ]
    return ContentModel(
        blocks=blocks,
        header="",
        footer=None,
        page_numbers=True,
        author="Kazma",
        subject="Letterhead B — English chrome proof",
    )


def _write_b_sample(*, logo: Path, font_dir: Path, language: str, stem: str, model: ContentModel) -> None:
    rtl = language.lower().startswith("ar")
    seed = "كاظمه" if rtl else "Kazma"
    profile = DocProfile.for_content(seed, language=language, rtl=rtl)
    profile.theme = dict(THEME)
    profile.theme["font_arabic"] = FACE
    profile.theme["font_latin"] = FACE
    profile.theme["body_size_ar"] = 12.0
    profile.theme["line_height_ar"] = 1.65

    docx = HERE / f"{stem}.docx"
    pdf = HERE / f"{stem}.pdf"
    try:
        DocxEngine(profile).render(model, docx)
    except PermissionError:
        stem = f"{stem}-reflow"
        docx = HERE / f"{stem}.docx"
        pdf = HERE / f"{stem}.pdf"
        DocxEngine(profile).render(model, docx)

    stamp = _format_stamp(datetime.now(), rtl=rtl)
    document = Document(str(docx))
    _letterhead_b(document, logo, rtl=rtl, chrome=profile.chrome, stamp=stamp)
    for section in document.sections:
        section.top_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Cm(0.55)
        section.different_first_page_header_footer = True
    last_chrome = None
    for child in document.element.body:
        if child.tag == qn("w:tbl"):
            last_chrome = child
            continue
        break
    if last_chrome is not None:
        _spacer_after(document, last_chrome, pt=12)

    section = document.sections[0]
    _empty_header(section.first_page_header)
    _add_running_header(section.header, rtl=rtl, chrome=profile.chrome, stamp=stamp)
    _copy_element_children(section.footer, section.first_page_footer)
    document.save(str(docx))

    _docx_to_pdf(docx, pdf, font_dir)
    _raster_all(pdf)
    print(stem, "lang=", language, "pdf-fonts=", _pdf_fonts(pdf))


def main() -> None:
    if not COPY.is_file():
        raise SystemExit(f"missing operator copy: {COPY}")
    if not LOGO_SRC.is_file():
        raise SystemExit(f"missing dark logo: {LOGO_SRC}")
    font_dir = _stage_plex()
    _install_user_fonts(font_dir)
    os.environ["KAZMA_DOCUMENT_FONT_DIR"] = str(font_dir)
    logo = _prepare_logo()
    with _patched_theme(
        font_arabic=FACE,
        font_latin=FACE,
        body_size_ar=12.0,
        line_height_ar=1.65,
    ):
        _write_b_sample(
            logo=logo, font_dir=font_dir, language="ar",
            stem="letterhead-B-long", model=_model_long(),
        )
        _write_b_sample(
            logo=logo, font_dir=font_dir, language="en",
            stem="letterhead-B-long-en", model=_model_long_en(),
        )


if __name__ == "__main__":
    main()
