"""DOCX engine for the unified document layer.

Consumes a :class:`~kazma_core.documents.content_model.ContentModel` under a
:class:`~kazma_core.documents.profile.DocProfile` and writes a ``.docx`` whose
direction + alignment are *correct by construction*.

The Word BiDi alignment rule — that under ``w:bidi`` the reading-start edge is
the *physical right* and must be encoded as ``w:jc="start"`` (never
``"right"``) — is applied in exactly one place: :meth:`DocxEngine._set_paragraph`.
Every block (title bar, heading, body, list, TOC, citation, table cell,
header/footer) asks for an *intent* (``start`` / ``justify`` / ``end``) and the
profile maps it to the right ``w:jc`` value. There is no schema-invalid
``tcPr/w:jc`` anywhere, and no call site picks a raw ``RIGHT``.

This module owns all DOCX OOXML emission. Markdown parsing is reused from
:mod:`kazma_core.documents.rich_render` (format-agnostic). The old per-element
post-pass walk (``docx_apply_document_rtl``) is gone: the RTL foundation is
applied once up front and each block sets its own direction as it is created.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Literal

from kazma_core.documents.content_model import (
    Block,
    BodyBlock,
    CitationBlock,
    ContentModel,
    HeadingBlock,
    ImageBlock,
    SpacerBlock,
    TableBlock,
    TitleBlock,
    TOCBlock,
)
from kazma_core.documents.profile import DocProfile
from kazma_core.documents.rich_render import parse_rich_blocks, try_parse_pipe_table_blob

logger = logging.getLogger(__name__)

__all__ = ["DocxEngine"]

_Intent = Literal["start", "justify", "end"]


class DocxEngine:
    """Render a :class:`ContentModel` to ``.docx`` under a :class:`DocProfile`.

    Instantiate per document: ``DocxEngine(profile).render(model, output)``.
    """

    def __init__(self, profile: DocProfile) -> None:
        self.profile = profile
        self.theme = profile.theme

    # ------------------------------------------------------------------ #
    # public entry
    # ------------------------------------------------------------------ #
    def render(self, model: ContentModel, output: Path | str, *,
               assets_dir: Any = None) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        self._assets_dir = assets_dir
        self._last_heading_text = ""
        document = Document()
        self._apply_foundation(document)

        # updateFields: tell Word/LibreOffice to update fields (TOC, PAGE) on
        # open, so the TOC populates with entries + page numbers and the footer
        # shows the real page number without a manual "update field" step.
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            settings = document.settings.element
            if settings.find(qn("w:updateFields")) is None:
                uf = OxmlElement("w:updateFields")
                uf.set(qn("w:val"), "true")
                settings.append(uf)
        except Exception:
            logger.debug("[docx] updateFields setting failed", exc_info=True)

        # Page setup (matches the PDF margins, not Word's 1.25" default).
        # Page size comes from the shared theme so DOCX and PDF share one
        # geometry (default.docx ships US Letter; we normalise to the theme's
        # A4 so the DOCX-route PDF matches the reportlab PDF exactly).
        from docx.shared import Mm
        page_mm = self.theme.get("page_size_mm", (210.0, 297.0))
        for section in document.sections:
            section.page_width = Mm(float(page_mm[0]))
            section.page_height = Mm(float(page_mm[1]))
            section.top_margin = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        # Normal style: Latin + Arabic typefaces, justified, theme spacing.
        try:
            from kazma_core.documents.style_theme import theme_fonts

            fonts = theme_fonts(rtl=self.profile.rtl)
            normal = document.styles["Normal"]
            normal.font.name = fonts["latin"]
            # Latin size stays at body_size even in RTL docs. Arabic is
            # sized independently via w:szCs (see theme_cs_size) — setting
            # Normal.font.size to body_size_ar pumped mixed English.
            latin_pt = float(self.theme["body_size"])
            normal.font.size = Pt(latin_pt)
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                from kazma_core.documents.style_theme import theme_cs_size

                r_pr = document.styles["Normal"].element.get_or_add_rPr()
                r_fonts = r_pr.find(qn("w:rFonts"))
                if r_fonts is None:
                    r_fonts = OxmlElement("w:rFonts")
                    r_pr.insert(0, r_fonts)
                r_fonts.set(qn("w:ascii"), fonts["latin"])
                r_fonts.set(qn("w:hAnsi"), fonts["latin"])
                r_fonts.set(qn("w:cs"), fonts["cs"])
                self._set_sz_cs(r_pr, theme_cs_size(latin_pt))
            except Exception:
                logger.debug("[docx] Normal rFonts/szCs failed", exc_info=True)
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            normal.paragraph_format.space_after = Pt(8)
            leading = (
                float(self.theme.get("line_height_ar", 1.85))
                if self.profile.rtl
                else float(self.theme["line_height"])
            )
            normal.paragraph_format.line_spacing = leading
        except Exception:
            logger.debug("[docx] Normal style setup failed", exc_info=True)

        # Header / footer chrome.
        self._write_header_footer(document, model)

        # File core properties (title/author/subject/keywords) from the model.
        self._set_core_properties(document, model)

        # Body blocks.
        for block in model.blocks:
            try:
                self._render_block(document, block)
            except Exception:
                logger.debug("[docx] block render failed: %r", block, exc_info=True)

        document.save(str(output))

    def _set_core_properties(self, document: Any, model: ContentModel) -> None:
        """Populate the file's core properties (shown in Explorer/Finder/search)."""
        try:
            title = ""
            for b in model.blocks:
                if isinstance(b, TitleBlock) and b.level == 0:
                    title = b.text
                    break
            cp = document.core_properties
            cp.title = title or "Document"
            cp.author = model.author or "Kazma"
            if model.subject:
                cp.subject = model.subject
            if model.keywords:
                cp.keywords = model.keywords
        except Exception:
            logger.debug("[docx] core properties failed", exc_info=True)

    # ================================================================== #
    # foundation — applied ONCE, up front
    # ================================================================== #
    def _apply_foundation(self, document: Any) -> None:
        """Set section/settings/Normal-style defaults so the doc opens RTL.

        Per-paragraph ``w:bidi``/``w:jc`` and per-run ``w:rtl`` are still set
        by each block as it is created (belt-and-suspenders, and necessary
        because not every consumer honors style inheritance reliably). But the
        document-level chrome (section direction, theme language, numbering)
        must come from here.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if not self.profile.rtl:
            return  # default.docx is already a correct LTR foundation

        # 1) Section: bidi + rtlGutter (explicit val=1 — empty bools are
        #    ignored by some consumers).
        for section in document.sections:
            sect_pr = section._sectPr
            for tag in ("w:bidi", "w:rtlGutter"):
                el = sect_pr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    sect_pr.append(el)
                el.set(qn("w:val"), "1")

        # 2) settings.xml: document language is Arabic (not an en-US shell).
        settings = document.settings.element
        tfl = settings.find(qn("w:themeFontLang"))
        if tfl is None:
            tfl = OxmlElement("w:themeFontLang")
            settings.append(tfl)
        tfl.set(qn("w:val"), "ar-SA")
        tfl.set(qn("w:bidi"), "ar-SA")
        tfl.set(qn("w:eastAsia"), "ar-SA")

        # 3) Normal style: paragraph bidi + Arabic lang. We deliberately do
        #    NOT put w:rtl on the style's rPr — that would force RTL on Latin
        #    runs too. Per-run w:rtl is applied selectively by _mark_run.
        try:
            normal = document.styles["Normal"].element
            p_pr = normal.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                normal.append(p_pr)
            if p_pr.find(qn("w:bidi")) is None:
                bidi = OxmlElement("w:bidi")
                bidi.set(qn("w:val"), "1")
                p_pr.append(bidi)
            r_pr = normal.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                normal.append(r_pr)
            lang = r_pr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                r_pr.append(lang)
            lang.set(qn("w:bidi"), "ar-SA")
        except Exception:
            logger.debug("[docx] Normal style RTL foundation failed", exc_info=True)

        # 4) List numbering: RTL levels so List Number items show numbers on
        #    the reading-start (right) side.
        self._ensure_numbering_rtl(document)

    def _ensure_numbering_rtl(self, document: Any) -> None:
        """Flip every numbering level to RTL (lvlJc right + right-side indent).

        Only affects paragraphs that actually use the List Number style; the
        TOC is plain numbered text and is handled separately.
        """
        from docx.oxml.ns import qn

        try:
            numbering_part = document.part.numbering_part
        except Exception:
            return
        if numbering_part is None:
            return
        elm = numbering_part._element
        for lvl in elm.iterdescendants(tag=qn("w:lvl")):
            lvl_jc = lvl.find(qn("w:lvlJc"))
            if lvl_jc is not None and lvl_jc.get(qn("w:val")) == "left":
                lvl_jc.set(qn("w:val"), "right")
            p_pr = lvl.find(qn("w:pPr"))
            if p_pr is not None:
                ind = p_pr.find(qn("w:ind"))
                if ind is not None:
                    left_val = ind.get(qn("w:left"))
                    if left_val is not None:
                        ind.set(qn("w:right"), left_val)
                        if qn("w:left") in ind.attrib:
                            del ind.attrib[qn("w:left")]

    # ================================================================== #
    # paragraph + run direction — the alignment policy application point
    # ================================================================== #
    def _set_paragraph(self, p: Any, intent: _Intent) -> None:
        """Set paragraph direction + alignment from an *intent*.

        This is the single place that maps intent → ``w:jc``. Under RTL it also
        stamps ``w:bidi`` and selectively marks Arabic runs. No call site may
        set ``w:jc`` or paragraph alignment directly.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Direction is resolved per PARAGRAPH, not per document. A mixed
        # document (a bilingual archive, a report quoting Arabic sources) has
        # English paragraphs and Arabic paragraphs, and each needs its own
        # ``w:bidi``. Taking it from the document meant whichever language lost
        # the ratio had every one of its paragraphs aligned backwards.
        #
        # The LTR branch stamps ``w:bidi val="0"`` explicitly rather than
        # leaving the element out: in an RTL document the Normal style and the
        # section both carry bidi, so an English paragraph that says nothing
        # inherits RTL.
        p_pr = p._p.get_or_add_pPr()
        block_rtl = self.profile.block_is_rtl(p.text or "")
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)
        bidi.set(qn("w:val"), "1" if block_rtl else "0")

        jc_val = self.profile.docx_jc(intent)
        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            p_pr.append(jc)
        jc.set(qn("w:val"), jc_val)
        if block_rtl:
            for run in p.runs:
                self._mark_run(run)

    def _mark_run(self, run: Any) -> None:
        """Mark a run as RTL complex-script — only if it contains Arabic.

        Latin runs are left neutral so Word's BiDi algorithm handles them
        naturally (no forced-reversal of embedded English/numbers).
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if not self.profile.run_is_rtl(run.text or ""):
            return
        r_pr = run._r.get_or_add_rPr()
        if r_pr.find(qn("w:rtl")) is None:
            r_pr.append(OxmlElement("w:rtl"))
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        from kazma_core.documents.style_theme import theme_fonts

        fonts = theme_fonts(rtl=True)
        for attr, val in (
            ("w:ascii", fonts["latin"]),
            ("w:hAnsi", fonts["latin"]),
            ("w:cs", fonts["arabic"]),
        ):
            r_fonts.set(qn(attr), val)
        lang = r_pr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            r_pr.append(lang)
        lang.set(qn("w:bidi"), "ar-SA")
        lang.set(qn("w:val"), "ar-SA")
        # Complex-script size is independent of Latin w:sz. Copying sz→szCs
        # made Sakkal match Calibri in nominal pt (optically smaller) and
        # body runs often have no per-run w:sz (they inherit Normal), so
        # they never got szCs and fell back to ~11pt. Always write szCs.
        # Bold/italic still need the Cs siblings (python-docx writes only
        # w:b/w:i). Insert after the Latin sibling for CT_RPr order.
        from kazma_core.documents.style_theme import theme_cs_size

        sz = r_pr.find(qn("w:sz"))
        latin_pt: float | None = None
        if sz is not None:
            try:
                half = int(sz.get(qn("w:val")) or 0)
                if half > 0:
                    latin_pt = half / 2.0
            except (TypeError, ValueError):
                latin_pt = None
        self._set_sz_cs(r_pr, theme_cs_size(latin_pt), after=sz)
        b = r_pr.find(qn("w:b"))
        if b is not None and r_pr.find(qn("w:bCs")) is None:
            b.addnext(OxmlElement("w:bCs"))
        i = r_pr.find(qn("w:i"))
        if i is not None and r_pr.find(qn("w:iCs")) is None:
            i.addnext(OxmlElement("w:iCs"))

    def _is_repeat_heading(self, text: str) -> bool:
        from kazma_core.documents.heading_text import headings_equivalent

        prev = getattr(self, "_last_heading_text", "") or ""
        return bool(text) and headings_equivalent(text, prev)

    @staticmethod
    def _keep_with_following(p: Any) -> None:
        """Pin a heading to the next paragraph (orphan-heading guard)."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_pr = p._p.get_or_add_pPr()
        for tag in ("w:keepNext", "w:keepLines", "w:widowControl"):
            el = p_pr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                p_pr.append(el)
            el.set(qn("w:val"), "1")

    @staticmethod
    def _force_ltr_paragraph(p: Any) -> None:
        """Pin a paragraph to LTR (code / math). Do not mark runs RTL."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pr = p._p.get_or_add_pPr()
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)
        bidi.set(qn("w:val"), "0")
        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            p_pr.append(jc)
        jc.set(qn("w:val"), "left")

    @staticmethod
    def _mark_ltr_run(run: Any) -> None:
        """Force a run to Latin LTR (no w:rtl) so source/math is not reversed."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        r_pr = run._r.get_or_add_rPr()
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is not None:
            r_pr.remove(rtl)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        face = run.font.name or "Consolas"
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            r_fonts.set(qn(attr), face)

    def _write_math(self, document: Any, tex: str, *, display: bool) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        from kazma_core.documents.math_text import latex_to_unicode

        shown = latex_to_unicode(tex)
        p = document.add_paragraph()
        run = p.add_run(shown)
        run.font.name = "Cambria Math"
        run.italic = True
        run.font.size = Pt(13 if display else 11)
        self._mark_ltr_run(run)
        self._force_ltr_paragraph(p)
        if display:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)

    def _write_code_block(self, document: Any, text: str) -> None:
        """LTR isolated code: one line per paragraph inside a single cell.

        An RTL section bidi-reverses a single wrapped run (the 'digest()'
        / comment-jumble symptom). Per-line LTR paragraphs in a table cell
        keep source order.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        table = document.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        fill = (self.theme.get("code_bg") or "#eff6ff").lstrip("#").upper()
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)
        lines = (text or "").splitlines() or [""]
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            self._mark_ltr_run(run)
            self._force_ltr_paragraph(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
        # Do not mark the table bidiVisual — code stays LTR even in RTL docs.
        self._decorate_table_paging(table, header=False, keep_together=True)
        sp = document.add_paragraph("")
        self._set_paragraph(sp, "start")

    def _decorate_table_paging(
        self,
        table: Any,
        *,
        header: bool = False,
        keep_together: bool | None = None,
    ) -> None:
        """Don't split rows; optionally repeat the header; keep small tables intact."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        nrows = len(table.rows)
        if keep_together is None:
            keep_together = nrows <= 12
        for i, row in enumerate(table.rows):
            tr = row._tr
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = OxmlElement("w:trPr")
                tr.insert(0, tr_pr)
            cant = tr_pr.find(qn("w:cantSplit"))
            if cant is None:
                cant = OxmlElement("w:cantSplit")
                tr_pr.append(cant)
            cant.set(qn("w:val"), "1")
            if header and i == 0:
                hdr = tr_pr.find(qn("w:tblHeader"))
                if hdr is None:
                    hdr = OxmlElement("w:tblHeader")
                    tr_pr.append(hdr)
                hdr.set(qn("w:val"), "1")
            if keep_together and i < nrows - 1:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._keep_with_following(p)

    @staticmethod
    def _set_sz_cs(r_pr: Any, size_pt: float, *, after: Any = None) -> None:
        """Set or update ``w:szCs`` (half-points) on a run/style rPr."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        half = max(1, int(round(float(size_pt) * 2)))
        szcs = r_pr.find(qn("w:szCs"))
        if szcs is None:
            szcs = OxmlElement("w:szCs")
            if after is not None:
                after.addnext(szcs)
            else:
                sz = r_pr.find(qn("w:sz"))
                if sz is not None:
                    sz.addnext(szcs)
                else:
                    r_pr.append(szcs)
        szcs.set(qn("w:val"), str(half))

    def _mark_table_rtl(self, table: Any) -> None:
        """``w:bidiVisual`` on a table so columns read right→left."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        for child in list(tbl_pr):
            if child.tag == qn("w:bidiVisual"):
                tbl_pr.remove(child)
        bidi_vis = OxmlElement("w:bidiVisual")
        bidi_vis.set(qn("w:val"), "1")
        tbl_pr.append(bidi_vis)

    # ================================================================== #
    # block dispatch
    # ================================================================== #
    def _render_block(self, document: Any, block: Block) -> None:
        if isinstance(block, TitleBlock):
            self._heading_bar(document, block.text, level=block.level, fill_hex=block.fill)
        elif isinstance(block, HeadingBlock):
            # Section headings are TOC-indexable (carry w:outlineLvl).
            self._heading_bar(document, block.text, level=block.level,
                              fill_hex=block.fill, indexable=True)
        elif isinstance(block, BodyBlock):
            self._write_rich_body(document, block.text)
        elif isinstance(block, TableBlock):
            if block.heading:
                self._heading_bar(document, block.heading, level=block.heading_level,
                                  fill_hex=self.theme["heading_fill"].lstrip("#"))
            self._add_table(document, block.headers, block.rows)
        elif isinstance(block, TOCBlock):
            self._write_toc(document, block.entries)
        elif isinstance(block, CitationBlock):
            self._write_citations(document, block.items)
        elif isinstance(block, ImageBlock):
            self._render_image(document, block)
        elif isinstance(block, SpacerBlock):
            p = document.add_paragraph(block.text or "")
            self._set_paragraph(p, "start")

    def _render_image(self, document: Any, block: ImageBlock) -> None:
        """Embed an approved-asset image (validated by the worker's _verify_assets).

        Only files present in the sha256-checked assets dir are reachable; a
        referenced-but-missing image is skipped (never embedded).
        """
        from docx.shared import Inches, Pt

        path = self._resolve_asset(block.name)
        if path is None:
            logger.debug("[docx] image not in approved assets, skipped: %s", block.name)
            return
        try:
            document.add_picture(str(path), width=Inches(block.width_in))
            if block.caption:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                cap = document.add_paragraph(block.caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                if self.profile.rtl:
                    p_pr = cap._p.get_or_add_pPr()
                    if p_pr.find(qn("w:bidi")) is None:
                        b = OxmlElement("w:bidi")
                        b.set(qn("w:val"), "1")
                        p_pr.append(b)
            document.add_paragraph("")  # spacer
        except Exception:
            logger.debug("[docx] image embed failed: %s", block.name, exc_info=True)

    def _resolve_asset(self, name: str) -> Path | None:
        if not name or not getattr(self, "_assets_dir", None):
            return None
        candidate = Path(str(self._assets_dir)) / Path(str(name)).name
        return candidate if candidate.is_file() else None

    # ================================================================== #
    # header / footer
    # ================================================================== #
    def _write_header_footer(self, document: Any, model: ContentModel) -> None:
        from docx.shared import Pt

        for section in document.sections:
            # Header: brand chrome.
            hp = section.header.paragraphs[0]
            hp.text = model.header or ""
            if (model.header or "").strip():
                for r in hp.runs:
                    r.font.size = Pt(8)
                self._set_paragraph(hp, "start")
            # Footer: brand + localized page label wrapping a PAGE field
            # ("صفحة {n}" / "Page {n}"). A bare digit in an Arabic footer
            # reads as leftover English chrome.
            fp = section.footer.paragraphs[0]
            fp.text = ""
            brand = model.footer or self.profile.chrome["brand"]
            brand_run = fp.add_run(brand)
            brand_run.font.size = Pt(8)
            if model.page_numbers:
                sep = fp.add_run("    —    ")
                sep.font.size = Pt(8)
                page_fmt = str(self.profile.chrome.get("page_fmt") or "{n}")
                prefix, _marker, suffix = page_fmt.partition("{n}")
                if prefix:
                    pre = fp.add_run(prefix)
                    pre.font.size = Pt(8)
                self._append_field(fp, "PAGE", "1")  # updated by Word/LibreOffice
                if suffix:
                    post = fp.add_run(suffix)
                    post.font.size = Pt(8)
            for r in fp.runs:
                if r.font.size is None:
                    r.font.size = Pt(8)
            if (fp.text or "").strip() or model.page_numbers:
                self._set_paragraph(fp, "start")

    @staticmethod
    def _append_field(paragraph: Any, instr: str, cache_text: str | list[str] = "") -> None:
        """Append a complex field (begin/instr/separate/result/end) to a paragraph.

        Used for the PAGE field (footer page numbers) and the TOC field. The
        ``cache_text`` is the pre-update display (Word/LibreOffice recompute the
        real value on open / update). A list renders as multiple lines (runs
        joined by ``<w:br/>``) so the TOC cache looks like a real list.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _fldchar(t: str) -> Any:
            r = OxmlElement("w:r")
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), t)
            r.append(fc)
            return r

        def _instr(text: str) -> Any:
            r = OxmlElement("w:r")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = f" {text} "
            r.append(it)
            return r

        def _text(text: str) -> Any:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            return r

        def _br() -> Any:
            r = OxmlElement("w:r")
            r.append(OxmlElement("w:br"))
            return r

        paragraph._p.append(_fldchar("begin"))
        paragraph._p.append(_instr(instr))
        paragraph._p.append(_fldchar("separate"))
        if cache_text:
            lines = cache_text if isinstance(cache_text, list) else [cache_text]
            for i, ln in enumerate(lines):
                if i > 0:
                    paragraph._p.append(_br())
                paragraph._p.append(_text(ln))
        paragraph._p.append(_fldchar("end"))

    # ================================================================== #
    # heading bar (full-width filled single-cell table)
    # ================================================================== #
    def _heading_bar(self, document: Any, text: str, *, level: int = 1,
                     fill_hex: str | None = None, indexable: bool = False) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        from kazma_core.documents.style_theme import theme_fonts

        fonts = theme_fonts(rtl=self.profile.rtl)
        sizes = {
            0: float(self.theme["title_size"]),
            1: float(self.theme["h1_size"]),
            2: float(self.theme["h2_size"]),
            3: float(self.theme["h3_size"]),
        }
        size = sizes.get(int(level), 12.0)
        ink = (self.theme.get("heading") or "#1e3a5f").lstrip("#")
        rule = (self.theme.get("accent") or "#3b82f6").lstrip("#").upper()

        if self._is_repeat_heading(text):
            return

        p = document.add_paragraph()
        run = p.add_run(text or "")
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = fonts["latin"]
        run.font.color.rgb = RGBColor(int(ink[0:2], 16), int(ink[2:4], 16), int(ink[4:6], 16))

        # Editorial: navy type + royal accent rule (no inverted bar).
        # Title gets a bottom rule; section headings also get a start-edge bar.
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "16" if int(level) == 0 else "8")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), rule)
        p_bdr.append(bottom)
        if int(level) in (1, 2):
            edge = "right" if self.profile.rtl else "left"
            side = OxmlElement(f"w:{edge}")
            side.set(qn("w:val"), "single")
            side.set(qn("w:sz"), "18")
            side.set(qn("w:space"), "8")
            side.set(qn("w:color"), rule)
            p_bdr.append(side)
        p_pr.append(p_bdr)

        self._set_paragraph(p, "start")
        # Keep the heading with the following body. Do NOT page-break-before
        # every heading — only hop to the next page when the heading would
        # otherwise sit alone at the bottom. An empty spacer paragraph here
        # would eat keepNext and re-orphan the body.
        self._keep_with_following(p)
        p.paragraph_format.space_before = Pt(10 if int(level) else 4)
        p.paragraph_format.space_after = Pt(8)
        if indexable and level >= 1:
            p_pr = p._p.get_or_add_pPr()
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), str(min(level, 3) - 1))
            p_pr.append(outline)
        self._last_heading_text = text or ""

    # ================================================================== #
    # table of contents
    # ================================================================== #
    def _write_toc(self, document: Any, entries: list[str]) -> None:
        """Insert a real Word TOC field (auto-updated by Word/LibreOffice).

        Section heading bars carry ``w:outlineLvl`` (see :meth:`_heading_bar`),
        so the TOC field's ``\\o "1-3"`` switch indexes them and regenerates the
        table — with page numbers — when the field is updated (Word: right-click
        → Update Field; LibreOffice updates on open / PDF export). The entries
        are pre-populated as the field's cached result so the TOC is never empty
        before the first update.
        """
        self._heading_bar(document, self.profile.chrome["toc"], level=2,
                          fill_hex=self.theme["heading_fill"].lstrip("#"))
        clean = [e for e in entries if e]
        if not clean:
            return
        cache = [f"{i}. {e}" for i, e in enumerate(clean, 1)]
        p = document.add_paragraph()
        self._append_field(p, 'TOC \\o "1-3" \\h \\z \\u', cache_text=cache)
        self._set_paragraph(p, "start")

    # ================================================================== #
    # citations (numbered references)
    # ================================================================== #
    def _write_citations(self, document: Any, items: list[str]) -> None:
        self._heading_bar(document, self.profile.chrome["references"], level=2,
                          fill_hex=self.theme["heading_fill"].lstrip("#"))
        for item in items:
            try:
                p = document.add_paragraph(str(item), style="List Number")
            except KeyError:
                p = document.add_paragraph(str(item))
            self._set_paragraph(p, "start")

    # ================================================================== #
    # structured table
    # ================================================================== #
    def _add_table(self, document: Any, headers: list[str], rows: list[list[str]]) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        from kazma_core.documents.style_theme import theme_fonts

        if not headers:
            return
        ncols = len(headers)
        table = document.add_table(rows=1 + len(rows), cols=ncols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        def _shade(cell: Any, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), fill)
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)

        def _fill(cell: Any, text: str, *, header: bool) -> None:
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            # Latin cell size stays 10pt; Arabic is sized via w:szCs in _mark_run.
            run.font.size = Pt(10)
            run.font.name = theme_fonts(rtl=self.profile.rtl)["latin"]
            if header:
                run.bold = True
                fg = (self.theme.get("table_header_fg") or "#16223a").lstrip("#")
                run.font.color.rgb = RGBColor(int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
                _shade(cell, (self.theme.get("table_header_bg") or "#eff6ff").lstrip("#").upper())
            else:
                _shade(cell, (self.theme.get("table_row_bg") or "#f8fafc").lstrip("#").upper())
            self._set_paragraph(p, "start")

        for ci, h in enumerate(headers):
            _fill(table.rows[0].cells[ci], str(h), header=True)
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                val = row[ci] if ci < len(row) else ""
                _fill(table.rows[ri + 1].cells[ci], val, header=False)

        if self.profile.rtl:
            self._mark_table_rtl(table)

        self._decorate_table_paging(table, header=True)

        # Spacer paragraph after the table.
        sp = document.add_paragraph("")
        self._set_paragraph(sp, "start")

    # ================================================================== #
    # rich Markdown body
    # ================================================================== #
    def _write_rich_body(self, document: Any, body: str) -> None:
        """Render Markdown body to DOCX blocks with correct direction."""
        from docx.shared import Pt, RGBColor

        # Whole-body collapsed one-line table? → structured table.
        maybe = try_parse_pipe_table_blob(body)
        if maybe is not None and body.count("\n") < 2 and body.count("|") > 6:
            self._add_table(document, list(maybe.get("headers") or []),
                            list(maybe.get("rows") or []))
            return

        blocks = parse_rich_blocks(body)
        if not blocks and body.strip():
            blocks = [{"type": "paragraph", "text": p.strip()}
                      for p in body.split("\n\n") if p.strip()]

        for block in blocks:
            btype = block["type"]
            if btype == "heading":
                level = min(3, max(1, int(block.get("level") or 2)))
                fill = "1E3A5F" if level <= 2 else "334155"
                self._heading_bar(document, block.get("text") or "",
                                  level=level, fill_hex=fill, indexable=True)
            elif btype == "paragraph":
                # A collapsed table the block classifier missed?
                maybe2 = try_parse_pipe_table_blob(block.get("text") or "")
                if maybe2 is not None:
                    self._add_table(document, list(maybe2.get("headers") or []),
                                    list(maybe2.get("rows") or []))
                    continue
                p = document.add_paragraph()
                self._add_inline_md_runs(p, block.get("text") or "")
                self._set_paragraph(p, "justify")
            elif btype == "math":
                self._write_math(document, block.get("text") or "", display=True)
            elif btype == "quote":
                p = document.add_paragraph()
                run = p.add_run(block.get("text") or "")
                run.italic = True
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                self._set_paragraph(p, "justify")
                p.paragraph_format.left_indent = Pt(18)
            elif btype == "code":
                self._write_code_block(document, block.get("text") or "")
            elif btype == "list":
                style = "List Number" if block.get("ordered") else "List Bullet"
                for item in block.get("items") or []:
                    try:
                        p = document.add_paragraph(style=style)
                    except KeyError:
                        p = document.add_paragraph()
                    if p.runs:
                        p.runs[0].text = ""
                    self._add_inline_md_runs(p, item.get("text") or "")
                    self._set_paragraph(p, "start")
            elif btype == "table":
                self._add_table(document, list(block.get("headers") or []),
                                list(block.get("rows") or []))
            elif btype == "hr":
                p = document.add_paragraph("─" * 40)
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_inline_md_runs(self, paragraph: Any, text: str) -> None:
        """Add bold/italic/code/math/plain runs for inline Markdown."""
        from kazma_core.documents.math_text import (
            latex_to_unicode,
            split_inline_math,
        )

        pattern = re.compile(
            r"\*\*(.+?)\*\*|__(.+?)__|"
            r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|"
            r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)|"
            r"`([^`]+)`"
        )

        def _md_chunk(chunk: str) -> None:
            pos = 0
            for m in pattern.finditer(chunk):
                if m.start() > pos:
                    paragraph.add_run(chunk[pos:m.start()])
                if m.group(1) is not None or m.group(2) is not None:
                    run = paragraph.add_run(
                        m.group(1) if m.group(1) is not None else m.group(2)
                    )
                    run.bold = True
                elif m.group(3) is not None or m.group(4) is not None:
                    run = paragraph.add_run(
                        m.group(3) if m.group(3) is not None else m.group(4)
                    )
                    run.italic = True
                elif m.group(5) is not None:
                    run = paragraph.add_run(m.group(5))
                    run.font.name = "Consolas"
                    self._mark_ltr_run(run)
                pos = m.end()
            if pos < len(chunk):
                paragraph.add_run(chunk[pos:])

        for kind, chunk in split_inline_math(text or ""):
            if kind == "math":
                run = paragraph.add_run(latex_to_unicode(chunk))
                run.font.name = "Cambria Math"
                run.italic = True
                self._mark_ltr_run(run)
            elif kind == "money":
                run = paragraph.add_run("$" + chunk.strip())
                run.font.name = "Consolas"
                self._mark_ltr_run(run)
            else:
                _md_chunk(chunk)
        if not text:
            paragraph.add_run("")
