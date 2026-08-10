"""Rich document rendering helpers — markdown structure + Arabic PDF shaping.

ReportLab draws LTR and does **not** shape Arabic (joining forms) or apply the
Unicode BiDi algorithm. Without ``arabic_reshaper`` + ``python-bidi``, Arabic
PDF text appears as disconnected, often reversed glyphs.

Word (DOCX) *does* OpenType shaping; for DOCX we set RTL/bidi paragraph flags
and semantic styles instead of pre-shaping.
"""

from __future__ import annotations

import html
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

__all__ = [
    "arabic_ratio",
    "is_arabic_dominant",
    "shape_for_pdf",
    "parse_rich_blocks",
    "try_parse_pipe_table_blob",
    "inline_markdown_to_reportlab",
    "pdf_flowables_from_body",
    "docx_write_rich_body",
    "docx_set_rtl_paragraph",
    "docx_force_justify",
    "docx_set_paragraph_shading",
    "docx_heading_bar",
    "docx_add_table",
    "docx_apply_document_rtl",
    "docx_set_run_rtl",
]

_AR_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_UL_RE = re.compile(r"^(\s*)([-*•])\s+(.+)$")
_OL_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.+)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
# GFM-style table row: | a | b |
_TABLE_ROW_RE = re.compile(r"^\|.+\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:\-|]+\|$")


def arabic_ratio(text: str) -> float:
    if not text:
        return 0.0
    letters = [c for c in text if c.isalpha() or _AR_RE.match(c)]
    if not letters:
        return 0.0
    ar = sum(1 for c in letters if _AR_RE.match(c))
    return ar / len(letters)


def is_arabic_dominant(text: str, *, threshold: float = 0.35) -> bool:
    """True when enough Arabic letters are present to drive RTL layout."""
    if not text or not _AR_RE.search(text):
        return False
    return arabic_ratio(text) >= threshold or bool(_AR_RE.search(text[:200]))


def shape_for_pdf(text: str) -> str:
    """Reshape + BiDi-reorder for ReportLab (LTR drawing engine)."""
    if not text or not _AR_RE.search(text):
        return text
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
    except ImportError:
        logger.warning(
            "[rich_render] arabic_reshaper/python-bidi missing — Arabic PDF may look broken"
        )
        return text
    try:
        reshaper = arabic_reshaper.ArabicReshaper(
            configuration={
                "delete_harakat": False,
                "support_ligatures": True,
            }
        )
        reshaped = reshaper.reshape(text)
        base = "R" if is_arabic_dominant(text) else None
        return get_display(reshaped, base_dir=base)
    except Exception:
        logger.debug("[rich_render] shape_for_pdf failed", exc_info=True)
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display

            return get_display(arabic_reshaper.reshape(text))
        except Exception:
            return text


def _split_pipe_row(row: str) -> list[str]:
    inner = row.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def _is_sep_cell(cell: str) -> bool:
    c = (cell or "").strip().replace(" ", "")
    return bool(c) and bool(re.fullmatch(r":?-+:?", c))


def try_parse_pipe_table_blob(text: str) -> dict[str, Any] | None:
    """Parse a GFM table that may be multi-line OR collapsed onto one line.

    LLMs often emit: ``| h1 | h2 | |---|---| | r1 | r2 |`` as a single paragraph.
    Empty cells between header/sep/rows are row delimiters in collapsed form.
    """
    raw = (text or "").strip()
    if raw.count("|") < 4:
        return None

    # Multi-line GFM
    lines = [ln.strip() for ln in raw.replace("\r\n", "\n").split("\n") if ln.strip()]
    if len(lines) >= 2 and _TABLE_ROW_RE.match(lines[0]) and _TABLE_SEP_RE.match(lines[1]):
        headers = _split_pipe_row(lines[0])
        rows = [_split_pipe_row(ln) for ln in lines[2:] if _TABLE_ROW_RE.match(ln)]
        if headers and rows:
            return {"type": "table", "headers": headers, "rows": rows}

    cells = [c.strip() for c in raw.strip().strip("|").split("|")]
    if len(cells) < 3:
        return None
    try:
        first_sep = next(i for i, c in enumerate(cells) if _is_sep_cell(c))
    except StopIteration:
        return None
    last_sep = first_sep
    while last_sep + 1 < len(cells) and _is_sep_cell(cells[last_sep + 1]):
        last_sep += 1

    headers = list(cells[:first_sep])
    while headers and headers[-1] == "":
        headers.pop()
    # drop leading empties
    while headers and headers[0] == "":
        headers.pop(0)
    ncols = len(headers)
    if ncols < 1:
        return None

    data = cells[last_sep + 1 :]
    # empty string acts as row break; also pack by ncols
    rows: list[list[str]] = []
    cur: list[str] = []
    for c in data:
        if c == "":
            if cur:
                # pad/truncate to ncols
                if len(cur) < ncols:
                    cur = cur + [""] * (ncols - len(cur))
                rows.append(cur[:ncols])
                cur = []
            continue
        cur.append(c)
        if len(cur) == ncols:
            rows.append(cur)
            cur = []
    if cur:
        if len(cur) < ncols:
            cur = cur + [""] * (ncols - len(cur))
        rows.append(cur[:ncols])
    rows = [r for r in rows if any(x.strip() for x in r)]
    if not rows:
        return None
    return {"type": "table", "headers": headers, "rows": rows}


def parse_rich_blocks(body: str) -> list[dict[str, Any]]:
    """Parse lightweight markdown into typed blocks for PDF/DOCX emitters.

    Supported:
      - AT1–H6 (``#`` … ``######``)
      - unordered / ordered lists (``-`` / ``*`` / ``1.``)
      - fenced code blocks (````)
      - horizontal rules
      - paragraphs (blank-line separated); inner lines join with space
      - blockquotes (``>``)
    """
    if not (body or "").strip():
        return []

    lines = body.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    i = 0
    para_buf: list[str] = []
    list_buf: list[dict[str, Any]] | None = None
    list_kind: str | None = None

    def flush_para() -> None:
        nonlocal para_buf
        joined = "\n".join(ln.strip() for ln in para_buf if ln.strip()).strip()
        # Also try space-joined collapsed tables
        space_joined = " ".join(ln.strip() for ln in para_buf if ln.strip()).strip()
        para_buf = []
        if not joined:
            return
        parsed = try_parse_pipe_table_blob(joined) or try_parse_pipe_table_blob(space_joined)
        if parsed is not None:
            blocks.append(parsed)
            return
        # Prefer space-joined for normal paragraphs (original behavior)
        blocks.append({"type": "paragraph", "text": space_joined})

    def flush_list() -> None:
        nonlocal list_buf, list_kind
        if list_buf and list_kind:
            blocks.append({"type": "list", "ordered": list_kind == "ol", "items": list_buf})
        list_buf = None
        list_kind = None

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Fenced code
        if stripped.startswith("```"):
            flush_para()
            flush_list()
            lang = stripped[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # closing fence
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(code_lines)})
            continue

        if not stripped:
            flush_para()
            flush_list()
            i += 1
            continue

        # Markdown table (multi-line GFM or collapsed one-liner from LLMs)
        if "|" in stripped and (
            (i + 1 < len(lines) and _TABLE_SEP_RE.match(lines[i + 1].strip()))
            or _is_sep_cell(stripped.strip("|").split("|")[0] if False else "")
            or re.search(r"\|\s*:?-+:?\s*\|", stripped)
        ):
            # Gather consecutive pipe-ish lines into one blob
            flush_para()
            flush_list()
            blob_lines = [stripped]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt:
                    break
                if "|" in nxt or _TABLE_SEP_RE.match(nxt):
                    blob_lines.append(nxt)
                    i += 1
                    continue
                break
            blob = "\n".join(blob_lines)
            parsed = try_parse_pipe_table_blob(blob)
            if parsed is not None:
                blocks.append(parsed)
                continue
            # Not a table — fall through as paragraph text
            para_buf.extend(blob_lines)
            continue

        hm = _HEADING_RE.match(stripped)
        if hm:
            flush_para()
            flush_list()
            blocks.append(
                {
                    "type": "heading",
                    "level": min(6, len(hm.group(1))),
                    "text": hm.group(2).strip(),
                }
            )
            i += 1
            continue

        if _HR_RE.match(stripped):
            flush_para()
            flush_list()
            blocks.append({"type": "hr"})
            i += 1
            continue

        if stripped.startswith(">"):
            flush_para()
            flush_list()
            quote_lines = [stripped.lstrip("> ").strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip("> ").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(quote_lines)})
            continue

        um = _UL_RE.match(raw)
        om = _OL_RE.match(raw)
        if um or om:
            flush_para()
            kind = "ol" if om else "ul"
            m = om or um
            assert m is not None
            indent = len(m.group(1).replace("\t", "    ")) // 2
            item_text = m.group(3).strip()
            if list_kind != kind:
                flush_list()
                list_kind = kind
                list_buf = []
            assert list_buf is not None
            list_buf.append({"text": item_text, "level": indent})
            i += 1
            continue

        # Continuation of list item (indented line)
        if list_buf is not None and (raw.startswith("  ") or raw.startswith("\t")):
            list_buf[-1]["text"] = (list_buf[-1]["text"] + " " + stripped).strip()
            i += 1
            continue

        flush_list()
        para_buf.append(stripped)
        i += 1

    flush_para()
    flush_list()
    return blocks


def inline_markdown_to_reportlab(text: str, *, shape_arabic: bool = True) -> str:
    """Convert a subset of inline markdown to ReportLab ``<para>`` mini-HTML.

    ReportLab Paragraph supports ``<b>``, ``<i>``, ``<u>``, ``<font>``,
    ``<br/>``, ``<link>``. We escape everything else first, then re-inject tags.
    """
    if not text:
        return ""

    # Tokenize inline patterns on the *raw* string, escape segments
    parts: list[str] = []
    pos = 0
    # Simpler sequential pass
    pattern = re.compile(
        r"\*\*(.+?)\*\*|__(.+?)__|"
        r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|"
        r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)|"
        r"`([^`]+)`|"
        r"\[([^\]]+)\]\(([^)]+)\)"
    )

    def _esc_shape(s: str) -> str:
        s2 = shape_for_pdf(s) if shape_arabic else s
        return html.escape(s2)

    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append(_esc_shape(text[pos : m.start()]))
        if m.group(1) is not None or m.group(2) is not None:
            inner = m.group(1) if m.group(1) is not None else m.group(2)
            parts.append(f"<b>{_esc_shape(inner)}</b>")
        elif m.group(3) is not None or m.group(4) is not None:
            inner = m.group(3) if m.group(3) is not None else m.group(4)
            parts.append(f"<i>{_esc_shape(inner)}</i>")
        elif m.group(5) is not None:
            parts.append(
                f'<font face="Courier" size="9" color="#0f172a">'
                f"{_esc_shape(m.group(5))}</font>"
            )
        elif m.group(6) is not None:
            label = _esc_shape(m.group(6))
            href = html.escape(m.group(7), quote=True)
            parts.append(f'<link href="{href}" color="#2563eb"><u>{label}</u></link>')
        pos = m.end()
    if pos < len(text):
        parts.append(_esc_shape(text[pos:]))
    return "".join(parts) if parts else _esc_shape(text)


def pdf_flowables_from_body(
    body: str,
    *,
    styles: dict[str, Any],
    shape_arabic: bool = True,
    Spacer: Any = None,
    Paragraph: Any = None,
    HRFlowable: Any = None,
    KeepTogether: Any = None,
    colors: Any = None,
    Table: Any = None,
    TableStyle: Any = None,
    font_name: str = "Helvetica",
    bold_font_name: str = "Helvetica-Bold",
) -> list[Any]:
    """Build ReportLab flowables from a rich body string."""
    if Paragraph is None or Spacer is None:
        from reportlab.platypus import Paragraph as _P
        from reportlab.platypus import Spacer as _S

        Paragraph = _P
        Spacer = _S
    if HRFlowable is None:
        try:
            from reportlab.platypus import HRFlowable as _HR
        except ImportError:
            _HR = None  # type: ignore[misc, assignment]
        HRFlowable = _HR
    if Table is None or TableStyle is None:
        try:
            from reportlab.platypus import Table as _T
            from reportlab.platypus import TableStyle as _TS

            Table, TableStyle = _T, _TS
        except ImportError:
            Table = TableStyle = None  # type: ignore[misc, assignment]
    if colors is None:
        try:
            from reportlab.lib import colors as _colors

            colors = _colors
        except ImportError:
            pass

    flow: list[Any] = []
    blocks = parse_rich_blocks(body)
    if not blocks and body.strip():
        blocks = [
            {"type": "paragraph", "text": p.strip()}
            for p in body.split("\n\n")
            if p.strip()
        ]

    body_style = styles["body"]
    quote_style = styles.get("quote", body_style)
    code_style = styles.get("code", body_style)
    bullet_style = styles.get("bullet", body_style)
    number_style = styles.get("number", body_style)
    heading_fill = styles.get("heading_fill")  # reportlab Color or None

    def _heading_bar(text_html: str, st: Any, level: int) -> None:
        """Heading with optional solid background bar (parity with styled PDF)."""
        para = Paragraph(text_html, st)
        if Table is None or colors is None or heading_fill is None:
            flow.append(para)
            flow.append(Spacer(1, 6))
            return
        fill = heading_fill if level <= 2 else colors.HexColor("#e2e8f0")
        tbl = Table([[para]], colWidths=["*"])
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), fill),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        flow.append(tbl)
        flow.append(Spacer(1, 8))

    def _md_table(headers: list[str], rows: list[list[str]]) -> None:
        if Table is None or TableStyle is None or colors is None:
            # Fallback: plain lines
            line = " | ".join(headers)
            flow.append(
                Paragraph(
                    inline_markdown_to_reportlab(line, shape_arabic=shape_arabic),
                    body_style,
                )
            )
            return

        def cell(val: str) -> str:
            return shape_for_pdf(val) if shape_arabic else val

        data = [[cell(h) for h in headers]]
        for row in rows:
            data.append([cell(c) for c in row])
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTNAME", (0, 0), (-1, 0), bold_font_name),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f8fafc")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        flow.extend((tbl, Spacer(1, 10)))

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            level = int(block.get("level") or 2)
            key = f"h{min(level, 3)}"
            st = styles.get(key, styles.get("h2", body_style))
            text_html = inline_markdown_to_reportlab(
                block["text"], shape_arabic=shape_arabic
            )
            _heading_bar(text_html, st, level)
        elif btype == "paragraph":
            text_html = inline_markdown_to_reportlab(
                block["text"], shape_arabic=shape_arabic
            )
            flow.append(Paragraph(text_html, body_style))
            flow.append(Spacer(1, 8))
        elif btype == "quote":
            text_html = inline_markdown_to_reportlab(
                block["text"], shape_arabic=shape_arabic
            )
            flow.append(Paragraph(f"<i>{text_html}</i>", quote_style))
            flow.append(Spacer(1, 8))
        elif btype == "code":
            raw = block.get("text") or ""
            escaped = html.escape(raw).replace("\n", "<br/>")
            flow.append(
                Paragraph(
                    f'<font face="Courier" size="8">{escaped}</font>', code_style
                )
            )
            flow.append(Spacer(1, 8))
        elif btype == "list":
            ordered = bool(block.get("ordered"))
            for idx, item in enumerate(block.get("items") or [], 1):
                level = int(item.get("level") or 0)
                prefix = f"{idx}." if ordered else "•"
                inner = inline_markdown_to_reportlab(
                    item.get("text") or "", shape_arabic=shape_arabic
                )
                line = f"{html.escape(prefix)}&nbsp;&nbsp;{inner}"
                st = number_style if ordered else bullet_style
                pad = "&nbsp;" * (level * 4)
                flow.append(Paragraph(pad + line, st))
            flow.append(Spacer(1, 6))
        elif btype == "table":
            _md_table(
                list(block.get("headers") or []),
                list(block.get("rows") or []),
            )
        elif btype == "hr":
            if HRFlowable is not None and colors is not None:
                flow.append(
                    HRFlowable(
                        width="100%",
                        thickness=0.6,
                        color=colors.HexColor("#cbd5e1"),
                        spaceBefore=4,
                        spaceAfter=10,
                    )
                )
            else:
                flow.append(Spacer(1, 10))
    return flow


def docx_force_justify(paragraph: Any) -> None:
    """Force Word justification via both alignment enum and w:jc=both."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_pr = paragraph._p.get_or_add_pPr()
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "both")


def docx_set_paragraph_shading(paragraph: Any, fill_hex: str = "1E3A5F") -> None:
    """Solid background fill on a paragraph (heading bar)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fill = (fill_hex or "1E3A5F").lstrip("#").upper()
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def docx_heading_bar(
    document: Any,
    text: str,
    *,
    level: int = 1,
    rtl: bool = False,
    fill_hex: str = "1E3A5F",
) -> Any:
    """Full-width filled heading bar (single-cell table) — reliable in Word/Telegram viewers.

    Built-in Heading styles often ignore paragraph shading; a table cell fill
    matches the PDF bar look for both EN and AR.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    fill = (fill_hex or "1E3A5F").lstrip("#").upper()
    sizes = {0: 18, 1: 14, 2: 12, 3: 11}
    size = sizes.get(int(level), 12)

    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    # Cell background
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tc_pr.append(shd)

    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if rtl:
        docx_set_rtl_paragraph(p, justify=False)  # also stamps run w:rtl
        docx_set_table_rtl(table)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Spacing after bar
    sp = document.add_paragraph("")
    if rtl:
        docx_set_rtl_paragraph(sp, justify=False)
    return table


def docx_set_run_rtl(run: Any) -> None:
    """Mark a run as RTL complex-script (required for Word text direction).

    Paragraph ``w:bidi`` without per-run ``w:rtl`` often still paints as LTR
    in Word / Telegram / WPS.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_pr = run._r.get_or_add_rPr()
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))
    # Complex-script font + locale
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    # Prefer Arial (has Arabic glyphs on Windows); cs = complex script
    for attr, val in (
        ("w:ascii", "Arial"),
        ("w:hAnsi", "Arial"),
        ("w:cs", "Arial"),
    ):
        if not r_fonts.get(qn(attr)):
            r_fonts.set(qn(attr), val)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:bidi"), "ar-SA")
    lang.set(qn("w:val"), "ar-SA")


def docx_set_rtl_paragraph(paragraph: Any, *, justify: bool = True) -> None:
    """Mark a python-docx paragraph as RTL and optionally justify."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if justify:
        docx_force_justify(paragraph)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:bidi"))
    if existing is None:
        bidi = OxmlElement("w:bidi")
        # Explicit val=1 — some consumers ignore empty boolean elements
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
    else:
        existing.set(qn("w:val"), "1")
    # Every run in the paragraph must carry w:rtl
    for run in paragraph.runs:
        try:
            docx_set_run_rtl(run)
        except Exception:
            logger.debug("[rich_render] run rtl failed", exc_info=True)


def docx_set_table_rtl(table: Any) -> None:
    """Mark a table as RTL-visual so columns read right→left in Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    # Remove existing bidiVisual if any, then set
    for child in list(tbl_pr):
        if child.tag == qn("w:bidiVisual"):
            tbl_pr.remove(child)
    bidi_vis = OxmlElement("w:bidiVisual")
    bidi_vis.set(qn("w:val"), "1")
    tbl_pr.append(bidi_vis)


def docx_apply_document_rtl(document: Any) -> None:
    """Document-level RTL so Word opens as an RTL document (not LTR shell).

    Paragraph ``w:bidi`` alone leaves the page chrome / gutter / table flow LTR
    in many viewers (Telegram preview, Word status bar still says LTR).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 1) Every section: bidi + rtlGutter (explicit val=1)
    for section in document.sections:
        sect_pr = section._sectPr
        bidi = sect_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            sect_pr.append(bidi)
        bidi.set(qn("w:val"), "1")
        gutter = sect_pr.find(qn("w:rtlGutter"))
        if gutter is None:
            gutter = OxmlElement("w:rtlGutter")
            sect_pr.append(gutter)
        gutter.set(qn("w:val"), "1")

    # 2) Document language fully Arabic (not en-US shell + ar bidi only)
    settings = document.settings.element
    tfl = settings.find(qn("w:themeFontLang"))
    if tfl is None:
        tfl = OxmlElement("w:themeFontLang")
        settings.append(tfl)
    tfl.set(qn("w:val"), "ar-SA")
    tfl.set(qn("w:bidi"), "ar-SA")
    tfl.set(qn("w:eastAsia"), "ar-SA")

    # 3) All tables visual RTL
    for table in document.tables:
        try:
            docx_set_table_rtl(table)
        except Exception:
            logger.debug("[rich_render] table rtl failed", exc_info=True)

    # 4) Every paragraph + every run (spacers, headers, footers, cells)
    def _walk_paragraphs() -> Any:
        yield from document.paragraphs
        for section in document.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for para in _walk_paragraphs():
        try:
            p_pr = para._p.get_or_add_pPr()
            bidi = p_pr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                p_pr.append(bidi)
            bidi.set(qn("w:val"), "1")
            # Empty spacers → right; body keeps justify if already set
            if not (para.text or "").strip():
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif para.alignment is None:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                docx_set_run_rtl(run)
        except Exception:
            logger.debug("[rich_render] para rtl walk failed", exc_info=True)

    # 5) Normal + Title styles default bidi + run rtl defaults
    try:
        for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
            try:
                style = document.styles[style_name]
            except KeyError:
                continue
            style_el = style.element
            p_pr = style_el.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                style_el.append(p_pr)
            bidi = p_pr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                p_pr.append(bidi)
            bidi.set(qn("w:val"), "1")
            # style-level rPr rtl
            r_pr = style_el.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                style_el.append(r_pr)
            if r_pr.find(qn("w:rtl")) is None:
                r_pr.append(OxmlElement("w:rtl"))
            lang = r_pr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                r_pr.append(lang)
            lang.set(qn("w:bidi"), "ar-SA")
            lang.set(qn("w:val"), "ar-SA")
        document.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        logger.debug("[rich_render] style rtl failed", exc_info=True)


def docx_add_table(
    document: Any,
    headers: list[str],
    rows: list[list[str]],
    *,
    rtl: bool = False,
) -> None:
    """Insert a styled Word table (header band + grid)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    if not headers:
        return
    ncols = len(headers)
    table = document.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"

    def _shade_cell(cell: Any, fill: str) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    def _fill_cell(cell: Any, text: str, *, header: bool = False) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade_cell(cell, "1E3A5F")
        else:
            _shade_cell(cell, "F8FAFC")
        if rtl:
            docx_set_rtl_paragraph(p, justify=False)  # para bidi + run rtl
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for ci, h in enumerate(headers):
        _fill_cell(table.rows[0].cells[ci], str(h), header=True)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ""
            _fill_cell(table.rows[ri + 1].cells[ci], str(val), header=False)
    if rtl:
        docx_set_table_rtl(table)
    # Spacer paragraph after table
    sp = document.add_paragraph("")
    if rtl:
        docx_set_rtl_paragraph(sp, justify=False)


def _docx_add_runs_with_inline_md(paragraph: Any, text: str) -> None:
    """Add runs for bold/italic/code/plain to a paragraph (no Arabic pre-shape)."""
    pattern = re.compile(
        r"\*\*(.+?)\*\*|__(.+?)__|"
        r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|"
        r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)|"
        r"`([^`]+)`"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(1) is not None or m.group(2) is not None:
            run = paragraph.add_run(m.group(1) if m.group(1) is not None else m.group(2))
            run.bold = True
        elif m.group(3) is not None or m.group(4) is not None:
            run = paragraph.add_run(m.group(3) if m.group(3) is not None else m.group(4))
            run.italic = True
        elif m.group(5) is not None:
            run = paragraph.add_run(m.group(5))
            run.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if not text:
        paragraph.add_run("")


def docx_write_rich_body(
    document: Any,
    body: str,
    *,
    rtl: bool = False,
) -> None:
    """Write a rich markdown body into a python-docx Document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    blocks = parse_rich_blocks(body)
    if not blocks and body.strip():
        blocks = [
            {"type": "paragraph", "text": p.strip()}
            for p in body.split("\n\n")
            if p.strip()
        ]

    for block in blocks:
        btype = block["type"]
        if btype == "heading":
            level = min(3, max(1, int(block.get("level") or 2)))
            fill = "1E3A5F" if level <= 2 else "334155"
            # Prefer table-cell bar (visible in Word / Telegram) over Heading+shd
            docx_heading_bar(
                document,
                block.get("text") or "",
                level=level,
                rtl=rtl,
                fill_hex=fill,
            )
        elif btype == "paragraph":
            # May still be a collapsed table the block classifier missed
            maybe = try_parse_pipe_table_blob(block.get("text") or "")
            if maybe is not None:
                docx_add_table(
                    document,
                    list(maybe.get("headers") or []),
                    list(maybe.get("rows") or []),
                    rtl=rtl,
                )
                continue
            p = document.add_paragraph()
            _docx_add_runs_with_inline_md(p, block.get("text") or "")
            if rtl:
                docx_set_rtl_paragraph(p, justify=True)
            else:
                docx_force_justify(p)
        elif btype == "quote":
            p = document.add_paragraph()
            run = p.add_run(block.get("text") or "")
            run.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            if rtl:
                docx_set_rtl_paragraph(p, justify=True)
            else:
                docx_force_justify(p)
            p.paragraph_format.left_indent = Pt(18)
        elif btype == "code":
            p = document.add_paragraph()
            run = p.add_run(block.get("text") or "")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif btype == "list":
            style = "List Number" if block.get("ordered") else "List Bullet"
            for item in block.get("items") or []:
                p = document.add_paragraph(style=style)
                if p.runs:
                    p.runs[0].text = ""
                _docx_add_runs_with_inline_md(p, item.get("text") or "")
                if rtl:
                    docx_set_rtl_paragraph(p, justify=False)
        elif btype == "table":
            docx_add_table(
                document,
                list(block.get("headers") or []),
                list(block.get("rows") or []),
                rtl=rtl,
            )
        elif btype == "hr":
            p = document.add_paragraph("─" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
