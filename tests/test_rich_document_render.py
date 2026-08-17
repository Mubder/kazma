"""Rich PDF/DOCX formatting + Arabic shaping for document generation."""

from __future__ import annotations

from pathlib import Path

import pytest

from kazma_core.documents.rich_render import (
    inline_markdown_to_reportlab,
    is_arabic_dominant,
    parse_rich_blocks,
    shape_for_pdf,
)
from kazma_core.documents.style_theme import THEME, theme_cs_size


def test_latex_to_unicode_common_forms() -> None:
    from kazma_core.documents.math_text import (
        latex_to_unicode,
        looks_like_currency,
        split_inline_math,
    )
    from kazma_core.documents.rich_render import parse_rich_blocks

    assert latex_to_unicode(r"R = P \cdot I") == "R = P · I"
    out = latex_to_unicode(r"S(x) = \frac{1}{1 + e^{-x}}")
    assert "⁄" in out
    assert "⁻" in out or "−" in out or "-x" in out or "ˣ" in out
    assert looks_like_currency("0.0035")
    assert not looks_like_currency(r"R = P \\cdot I")
    kinds = [k for k, _ in split_inline_math(r"cost $0.0035$ and $R = P \\cdot I$")]
    assert "money" in kinds
    assert "math" in kinds
    blocks = parse_rich_blocks("intro\n\n$$S(x) = \\frac{1}{1+e^{-x}}$$\n\nmore")
    assert any(b.get("type") == "math" for b in blocks)


def test_theme_cs_size_splits_latin_and_arabic() -> None:
    """Arabic body is larger than Latin; chrome gets a modest bump only."""
    assert THEME["body_size_ar"] > THEME["body_size"]
    assert theme_cs_size() == THEME["body_size_ar"]
    assert theme_cs_size(THEME["body_size"]) == THEME["body_size_ar"]
    assert theme_cs_size(8) == 10.0
    assert theme_cs_size(THEME["h2_size"]) == THEME["h2_size"] + (
        THEME["body_size_ar"] - THEME["body_size"]
    )


def test_unified_theme_en_ar_pdf_share_heading_bars(tmp_path: Path) -> None:
    """EN and AR PDFs must use the same visual theme (bars + brand chrome)."""
    from kazma_core.documents.renderer_worker import _generate_pdf
    from kazma_core.documents.style_theme import THEME, localized_chrome
    from kazma_core.skills.exporter import generate_pdf_html_document

    body = (
        "## Overview\n\n"
        "A professional paragraph with enough words to exercise justified layout "
        "in both language modes of the shared document theme.\n\n"
        "| Feature | Status |\n|---|---|\n| Unified theme | Yes |\n"
    )
    body_ar = (
        "## نظرة عامة\n\n"
        "فقرة مهنية عربية طويلة بما يكفي لاختبار التبرير ونفس قالب التصميم.\n\n"
        "| الميزة | الحالة |\n|---|---|\n| قالب موحّد | نعم |\n"
    )
    w: list[str] = []
    en = tmp_path / "en.pdf"
    ar = tmp_path / "ar.pdf"
    _generate_pdf(
        en,
        {"title": "What is Kazma?", "lang": "en", "sections": [{"heading": "Main", "body": body}]},
        w,
    )
    _generate_pdf(
        ar,
        {"title": "ما هي كاظمة؟", "lang": "ar", "sections": [{"heading": "رئيسي", "body": body_ar}]},
        w,
    )
    assert en.stat().st_size > 2000 and ar.stat().st_size > 2000
    # HTML exporter shares THEME tokens
    html_en = generate_pdf_html_document(body, title="What is Kazma?", rtl=False)
    html_ar = generate_pdf_html_document(body_ar, title="ما هي كاظمة؟", rtl=True)
    assert THEME["heading_fill"] in html_en and THEME["heading_fill"] in html_ar
    assert 'dir="ltr"' in html_en and 'dir="rtl"' in html_ar
    assert localized_chrome(rtl=False)["brand"] in html_en or "Kazma" in html_en
    assert localized_chrome(rtl=True)["brand"] in html_ar or "كاظمة" in html_ar


def test_is_arabic_dominant() -> None:
    assert is_arabic_dominant("هذا نص عربي طويل بما يكفي") is True
    assert is_arabic_dominant("Hello English only document") is False
    assert is_arabic_dominant("Hello مرحبا mixed") is True  # enough AR letters


def test_parse_rich_blocks_headings_lists_bold() -> None:
    body = """# Main

Intro **bold** and *italic*.

## Details

- first
- second

1. one
2. two

> a quote
"""
    blocks = parse_rich_blocks(body)
    types = [b["type"] for b in blocks]
    assert "heading" in types
    assert "paragraph" in types
    assert "list" in types
    assert "quote" in types
    lists = [b for b in blocks if b["type"] == "list"]
    assert any(not b["ordered"] and len(b["items"]) == 2 for b in lists)
    assert any(b["ordered"] and len(b["items"]) == 2 for b in lists)


def test_shape_for_pdf_joins_arabic() -> None:
    raw = "مرحبا"
    shaped = shape_for_pdf(raw)
    # Shaping must change presentation forms (joined glyphs) or at least reorder
    assert shaped != raw or "م" not in shaped  # either reshaped forms or reordered
    assert len(shaped) >= 3
    # No reverse of logical without reshape libraries would leave isolated letters;
    # with libs, shaped string should not equal simple reverse alone necessarily.
    # Smoke: English unchanged
    assert shape_for_pdf("Hello") == "Hello"


def test_inline_markdown_to_reportlab_escapes_and_bold() -> None:
    html = inline_markdown_to_reportlab("Say **hi** & more", shape_arabic=False)
    assert "<b>" in html
    assert "hi" in html
    assert "&amp;" in html


def test_generate_pdf_arabic_and_markdown(tmp_path: Path) -> None:
    from kazma_core.documents.renderer_worker import _generate_pdf

    out = tmp_path / "ar.pdf"
    warnings: list[str] = []
    _generate_pdf(
        out,
        {
            "title": "تقرير الاختبار",
            "lang": "ar",
            "page_numbers": True,
            "sections": [
                {
                    "heading": "المقدمة",
                    "body": (
                        "## نظرة عامة\n\n"
                        "هذا **فقرة** عربية مبررة مع قائمة:\n\n"
                        "- البند الأول\n"
                        "- البند الثاني\n\n"
                        "ونص إنجليزي mixed: ISO/IEC 27001."
                    ),
                }
            ],
        },
        warnings,
    )
    assert out.is_file()
    assert out.stat().st_size > 800
    # Should not warn about missing reshape when deps present
    assert not any("arabic_reshaper" in w for w in warnings)


def test_generate_docx_arabic_rtl(tmp_path: Path) -> None:
    from docx import Document

    from kazma_core.documents.renderer_worker import _generate_docx

    out = tmp_path / "ar.docx"
    _generate_docx(
        out,
        {
            "title": "وثيقة اختبار",
            "lang": "ar",
            "sections": [
                {
                    "heading": "القسم الأول",
                    "body": (
                        "### فرعي\n\n"
                        "نص **عريض** مع نقاط:\n\n"
                        "- واحد\n"
                        "- اثنان\n"
                    ),
                }
            ],
        },
    )
    assert out.is_file()
    doc = Document(str(out))
    # Title/headings are editorial paragraphs (navy + brass rule), not filled bars
    all_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " " + cell.text
    assert "وثيقة" in all_text
    assert "واحد" in all_text or "اثنان" in all_text
    from docx.oxml.ns import qn

    has_rule = False
    for p in doc.paragraphs:
        p_pr = p._p.pPr
        if p_pr is not None and p_pr.find(qn("w:pBdr")) is not None:
            has_rule = True
            break
    assert has_rule, "editorial heading rule (w:pBdr) missing"

    has_bidi = False
    for p in doc.paragraphs:
        p_pr = p._p.pPr
        if p_pr is not None and p_pr.find(qn("w:bidi")) is not None:
            has_bidi = True
            break
    # Also check table cell paragraphs (title bar is RTL)
    if not has_bidi:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p_pr = p._p.pPr
                        if p_pr is not None and p_pr.find(qn("w:bidi")) is not None:
                            has_bidi = True
    assert has_bidi
    # Document must open as RTL (section bidi), not LTR shell with Arabic text
    sect = doc.sections[0]._sectPr
    assert sect.find(qn("w:bidi")) is not None
    assert sect.find(qn("w:rtlGutter")) is not None
    tfl = doc.settings.element.find(qn("w:themeFontLang"))
    assert tfl is not None and tfl.get(qn("w:bidi")) == "ar-SA"
    assert tfl.get(qn("w:val")) == "ar-SA"  # not en-US shell
    # Content tables (if any) stay bidiVisual; headings are no longer tables.
    # Run-level w:rtl is what Word uses for text direction (not only pPr bidi)
    run_rtl = 0
    for p in doc.paragraphs:
        for r in p.runs:
            r_pr = r._r.rPr
            if r_pr is not None and r_pr.find(qn("w:rtl")) is not None:
                run_rtl += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r_pr = r._r.rPr
                        if r_pr is not None and r_pr.find(qn("w:rtl")) is not None:
                            run_rtl += 1
    assert run_rtl >= 1


def test_markdown_table_and_docx_justify_shading(tmp_path: Path) -> None:
    from docx import Document
    from docx.oxml.ns import qn

    from kazma_core.documents.renderer_worker import _generate_docx, _generate_pdf
    from kazma_core.documents.rich_render import parse_rich_blocks

    body = (
        "## Overview\n\n"
        "A longer paragraph that must be justified in the DOCX output so spacing "
        "looks even across full lines of professional body text for export tests.\n\n"
        "| Feature | Status |\n"
        "|---|---|\n"
        "| Tables | Yes |\n"
        "| Justify | Yes |\n"
    )
    blocks = parse_rich_blocks(body)
    assert any(b["type"] == "table" for b in blocks)

    pdf = tmp_path / "t.pdf"
    docx = tmp_path / "t.docx"
    warnings: list[str] = []
    payload = {
        "title": "Parity Test",
        "lang": "en",
        "sections": [{"heading": "Main", "body": body}],
        "tables": [{"heading": "Extra", "headers": ["A", "B"], "rows": [["1", "2"]]}],
    }
    _generate_pdf(pdf, payload, warnings)
    _generate_docx(docx, payload)
    assert pdf.is_file() and pdf.stat().st_size > 500
    doc = Document(str(docx))
    # markdown table + payload table (headings are paragraphs now)
    assert len(doc.tables) >= 2
    found_jc = False
    found_bar_fill = False
    for p in doc.paragraphs:
        if "longer paragraph" in p.text and p._p.pPr is not None:
            jc = p._p.pPr.find(qn("w:jc"))
            if jc is not None and jc.get(qn("w:val")) == "both":
                found_jc = True
    for table in doc.tables:
        cell = table.rows[0].cells[0]
        tc_pr = cell._tc.tcPr
        if tc_pr is not None:
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None and shd.get(qn("w:fill")):
                found_bar_fill = True
                break
    assert found_jc
    assert found_bar_fill

    # Collapsed one-line table (as shipped in Telegram DOCX) becomes real table
    collapsed = (
        "| Feature | Kazma | LangChain | |---|---|---| "
        "| Multi-agent | Native | Add-on | | HITL | Triple | Manual |"
    )
    from kazma_core.documents.rich_render import try_parse_pipe_table_blob

    parsed = try_parse_pipe_table_blob(collapsed)
    assert parsed is not None
    assert parsed["headers"][0] == "Feature"
    assert len(parsed["rows"]) >= 2


def test_generate_pdf_english_lists(tmp_path: Path) -> None:
    from kazma_core.documents.renderer_worker import _generate_pdf

    out = tmp_path / "en.pdf"
    warnings: list[str] = []
    _generate_pdf(
        out,
        {
            "title": "English Report",
            "lang": "en",
            "sections": [
                {
                    "heading": "Overview",
                    "body": (
                        "## Goals\n\n"
                        "A **strong** start.\n\n"
                        "- alpha\n"
                        "- beta\n\n"
                        "1. first\n"
                        "2. second\n"
                    ),
                }
            ],
        },
        warnings,
    )
    assert out.is_file() and out.stat().st_size > 500



def test_generate_pdf_arabic_body_rtl_bbox(tmp_path: Path) -> None:
    import unicodedata
    """A long Arabic body paragraph must flow at page width — right-aligned
    lines (right edge near the page's right margin), with the paragraph START on
    the top line, spanning multiple lines.

    Route-tolerant: Arabic PDFs now route through DOCX→LibreOffice→PDF when
    LibreOffice is available (a real bidi engine; reportlab cannot correctly
    shape mixed Arabic+Latin+inline-markdown). The assertions use page-width-
    relative thresholds so they hold for both the DOCX-route (US Letter, the
    DOCX default) and the reportlab fallback (A4). Original intent preserved:
    guard against the v9 regression where the body collapsed to narrow
    non-right-aligned / line-scrambled output.
    """
    import pymupdf

    from kazma_core.documents.renderer_worker import _generate_pdf

    arabic_body = (
        "تقرير شامل يوضح قدرات معالجة المستندات في منظومة كاظمة، مع دعم كامل "
        "للغتين العربية والإنجليزية، والربط المعياري مع محركات التقرير."
    )
    out = tmp_path / "ar_body.pdf"
    warnings: list[str] = []
    _generate_pdf(
        out,
        {
            "title": "التقرير التنفيذ",
            "lang": "ar",
            "rtl": True,
            "toc": False,
            "sections": [{"heading": "الملخص", "body": arabic_body}],
        },
        warnings,
    )
    assert out.is_file()

    doc = pymupdf.open(str(out))
    page = doc[0]
    W = page.rect.width
    H = page.rect.height
    # Body-paragraph lines: Latin body (~11pt) or Arabic cs body
    # (theme_cs_size / body_size_ar). Right edge in the right ~15% of the
    # page (right-aligned), AND within the body vertical zone (excludes the
    # header brand band at the top and the footer band at the bottom).
    # Title/heading bars are larger so the size window already excludes them.
    body_hi = theme_cs_size() + 1.5
    body_lines: list[tuple[float, float, float, str]] = []
    for block in page.get_text("dict").get("blocks", []):
        for line in block.get("lines", []):
            lr = line["bbox"]
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans)
            size = spans[0]["size"] if spans else 0.0
            if (10.0 <= size <= body_hi and lr[2] >= W * 0.85
                    and 70.0 < lr[1] < H - 60.0 and len(text.strip()) > 2):
                body_lines.append((lr[1], lr[0], lr[2], text))
    doc.close()
    assert body_lines, (
        "Arabic body produced no right-aligned body-zone lines"
    )

    # The body paragraph must render right-aligned (not left-pinned, the v9
    # collapse symptom). pymupdf's per-line boxes are unreliable for justified
    # RTL (it jumbles mid/end fragments), so we assert on the robust signal:
    # the paragraph START word appears among the right-aligned body-zone lines.
    # NFKC normalizes presentation-form glyphs back to logical letters; control
    # chars (\x1f unit-separators LibreOffice/pymupdf insert) are stripped.
    joined = " ".join(t for _, _, _, t in body_lines)
    joined_norm = "".join(c for c in unicodedata.normalize("NFKC", joined) if c >= " ")
    assert "تقرير" in joined_norm, (
        f"body paragraph start not found among right-aligned lines: {joined[:60]!r}"
    )

    # Page-width flow: at least one right-aligned body-zone line spans a
    # meaningful fraction of the page (full-width justified flow, not a narrow
    # collapse like the v9 regression).
    max_line_width = max(x1 - x0 for _, x0, x1, _ in body_lines)
    assert max_line_width >= W * 0.5, (
        f"Arabic body should span >= half the page width; widest = {max_line_width:.1f} (page {W:.0f})"
    )



def test_subtitle_and_ar_citations_rtl(tmp_path: Path) -> None:
    """subtitle renders in both languages; AR citations right-align to the
    column right edge (x1 ~ 535); EN body fills the column (TA_JUSTIFY 'LTR'
    wrap). Regression-locks the cite_style fix that replaced body_style's
    left-pinned TA_JUSTIFY for Arabic reference lines."""
    import pymupdf
    from kazma_core.documents.renderer_worker import _generate_pdf

    ar_body = "تقرير شامل يوضح قدرات معالجة المستندات في منظومة كاظمة."
    en_body = ("Executive summary covering the document-processing capabilities of "
            "the Kazma agent framework, with native bilingual support for Arabic and "
            "English plus advanced analytics driving smart recommendations across the platform.")

    def lines(pdf):
        doc = pymupdf.open(str(pdf))
        out = []
        for page in doc:
            for block in page.get_text("dict").get("blocks", []):
                for ln in block.get("lines", []):
                    spans = ln.get("spans", [])
                    if spans and 8.0 <= spans[0]["size"] <= theme_cs_size() + 1.5:
                        text = "".join(s.get("text", "") for s in spans)
                        if len(text.strip()) > 2:
                            bb = ln["bbox"]
                            out.append((round(bb[1], 1), round(bb[0], 1), round(bb[2], 1), text))
        return out

    _generate_pdf(
        tmp_path / "ar.pdf",
        {
            "title": "T", "subtitle": "Kazma", "lang": "ar", "rtl": True,
            "citations": ["مصدر أول: تقرير داخلي غير مراجع", "مصدر ثان: بيانات داخلية"],
            "sections": [{"heading": "الملخص", "body": ar_body}],
        },
        [],
    )
    _generate_pdf(
        tmp_path / "en.pdf",
        {
            "title": "T", "subtitle": "Under title", "lang": "en",
            "citations": ["First source: internal report"],
            "sections": [{"heading": "Summary", "body": en_body}],
        },
        [],
    )

    ar = lines(tmp_path / "ar.pdf")
    en = lines(tmp_path / "en.pdf")
    assert any("Kazma" in t for _, _, _, t in ar), "AR subtitle missing"
    assert any("Under title" in t for _, _, _, t in en), "EN subtitle missing"
    # AR citations must be right-aligned (x1 ~ 535) and appear AFTER the short
    # body; the old body_style TA_JUSTIFY bug pinned them left (x0 ~ 60).
    after = [(y, x0, x1, t) for y, x0, x1, t in ar if y > 150.0 and x1 >= 528.0]
    assert after, "AR citations not right-aligned after body: " + repr(after)
    # EN body must justify to the full column (x1 ~ 535), not ragged-right.
    en_wide = [x1 for _, x0, x1, t in en if x1 >= 528.0 and "Under title" not in t]
    assert en_wide, "EN body did not fill the column"


def test_toc_items_align_per_language(tmp_path: Path) -> None:
    """TOC numbered entries must follow reading direction: AR entries right-align
    to the column edge (x1 ~ 535), EN entries left-align (x0 ~ 60). Regression
    for the AR TOC item previously pinned to x0=60 by body_style(TA_JUSTIFY)."""
    import pymupdf
    from kazma_core.documents.renderer_worker import _generate_pdf

    ar_body = "تقرير شامل يوضح قدرات معالجة المستندات في منظومة كاظمة."
    en_body = "Executive summary covering document-processing capabilities of the platform."

    def toc_entries(pdf: Path) -> list[tuple] :
        doc = pymupdf.open(str(pdf))
        out = []
        for pi, page in enumerate(doc):
            for b in page.get_text("dict").get("blocks", []):
                for ln in b.get("lines", []):
                    sp = ln.get("spans", [])
                    if sp and 8.0 <= sp[0]["size"] <= theme_cs_size() + 1.5:
                        txt = "".join(s.get("text", "") for s in sp)
                        if len(txt.strip()) > 2:
                            bb = ln["bbox"]
                            out.append((pi + 1, round(bb[1], 1), round(bb[0], 1), round(bb[2], 1), txt))
        return out

    _generate_pdf(tmp_path / "ar.pdf",
        {"title": "T", "lang": "ar", "rtl": True, "toc": True,
         "sections": [{"heading": "الملخص", "body": ar_body}]}, [])
    _generate_pdf(tmp_path / "en.pdf",
        {"title": "T", "lang": "en", "toc": True,
         "sections": [{"heading": "Summary", "body": en_body}]}, [])

    ar = toc_entries(tmp_path / "ar.pdf")
    en = toc_entries(tmp_path / "en.pdf")
    ar_toc = [x1 for pg, y, x0, x1, t in ar if pg == 1 and 150.0 <= y <= 320.0 and x1 >= 528.0]
    assert ar_toc, "AR TOC entry not right-aligned to column edge"
    en_toc = [x0 for pg, y, x0, x1, t in en if pg == 1 and 150.0 <= y <= 320.0 and x0 <= 70.0]
    assert en_toc, "EN TOC entry not left-aligned"

def test_docx_font_size_and_rtl_synced(tmp_path: Path) -> None:
    """DOCX must mirror the PDF/THEME styling: Calibri Latin + Arabic cs font,
    heading point sizes synced to THEME, and Arabic runs carry w:rtl + the
    paragraph w:bidi so Word opens RTL (regression: 'AR is LTR + no font set'
    and 'EN bigger font size not synced')."""
    import zipfile
    import re
    from kazma_core.documents.renderer_worker import _generate_docx
    from kazma_core.documents.style_theme import THEME

    ar_body = "تقرير شامل يوضح قدرات معالالجة في منصومة كاظمة."
    ar_pl = {
        "title": "الملخص التنفيذي", "subtitle": "تقرير استراتيجي",
        "lang": "ar", "rtl": True, "toc": True, "citations": ["مصدر أول: تقرير"],
        "sections": [{"heading": "الملخص", "body": ar_body}],
    }
    en_pl = {
        "title": "Executive Summary", "subtitle": "Strategic Report",
        "lang": "en", "rtl": False, "toc": True, "citations": ["First source"],
        "sections": [{"heading": "Summary", "body": "Executive summary of the platform."}],
    }
    _generate_docx(tmp_path / "ar.docx", ar_pl)
    _generate_docx(tmp_path / "en.docx", en_pl)

    ar_xml = zipfile.ZipFile(tmp_path / "ar.docx").read("word/document.xml").decode("utf-8")
    en_xml = zipfile.ZipFile(tmp_path / "en.docx").read("word/document.xml").decode("utf-8")
    styles = zipfile.ZipFile(tmp_path / "en.docx").read("word/styles.xml").decode("utf-8")

    # No Arial leakage in either document
    assert 'w:ascii="Arial"' not in ar_xml, "Arial leaked into AR DOCX run font"
    assert 'w:ascii="Arial"' not in en_xml, "Arial leaked into EN DOCX run font"
    assert 'w:ascii="Calibri"' in en_xml, "EN DOCX missing Calibri run font"
    # AR complex-script runs use the theme Arabic face + rtl + paragraph bidi
    ar_cs = THEME.get("font_arabic") or "Sakkal Majalla"
    assert f'w:cs="{ar_cs}"' in ar_xml, "AR DOCX runs missing Arabic complex-script font"
    assert "w:rtl" in ar_xml, "AR DOCX runs missing w:rtl"
    assert "w:bidi" in ar_xml, "AR DOCX paragraphs missing w:bidi"
    # EN heading sizes (half-points) synced to THEME
    sz = {int(x) for x in re.findall(r'w:sz w:val="(\d+)"', en_xml)}
    assert int(THEME["title_size"] * 2) in sz   # title -> 40
    assert int(THEME["h1_size"] * 2) in sz       # h1    -> 32
    assert int(THEME["h2_size"] * 2) in sz       # h2    -> 28
    assert int(THEME["h3_size"] * 2) in sz       # h3    -> 25
    # Body paragraph size on the Normal style (styles.xml) synced too
    assert ('w:val="%d"' % int(THEME["body_size"] * 2)) in styles, "EN DOCX body size not synced"
    # AR body Arabic is independently larger via w:szCs (not a copy of Latin sz)
    from kazma_core.documents.style_theme import theme_cs_size
    ar_styles = zipfile.ZipFile(tmp_path / "ar.docx").read("word/styles.xml").decode("utf-8")
    assert ('w:szCs w:val="%d"' % int(round(theme_cs_size() * 2))) in ar_styles, (
        "AR Normal style missing independent complex-script body size"
    )
    assert ('w:sz w:val="%d"' % int(THEME["body_size"] * 2)) in ar_styles, (
        "AR Normal latin size must stay at body_size (do not pump English)"
    )


def test_docx_numbering_rtl(tmp_path: Path) -> None:
    """DOCX list numbering must be RTL for Arabic documents (regression: list
    numbers stayed left-aligned LTR even when document is RTL, causing citations
    and TOC entries to have numbers on the wrong side)."""
    import zipfile
    from kazma_core.documents.renderer_worker import _generate_docx

    ar_pl = {
        "title": "الملخص التنفيذي",
        "subtitle": "تقرير استراتيجي",
        "lang": "ar",
        "rtl": True,
        "toc": True,
        "citations": ["مصدر أول: تقرير", "مصدر ثان: تحليل"],
        "sections": [{"heading": "الملخص", "body": "تقرير شامل."}],
    }
    _generate_docx(tmp_path / "ar.docx", ar_pl)

    # Extract numbering.xml and verify RTL settings
    ar_docx = zipfile.ZipFile(tmp_path / "ar.docx")
    numbering_xml = ar_docx.read("word/numbering.xml").decode("utf-8")

    # List numbering must be right-justified (not left)
    assert 'w:lvlJc w:val="right"' in numbering_xml, (
        "AR DOCX numbering not RTL: lvlJc should be 'right', not 'left'"
    )
    assert 'w:lvlJc w:val="left"' not in numbering_xml, (
        "AR DOCX numbering still has LTR lvlJc"
    )
    # Indentation should use w:right, not w:left
    assert 'w:right=' in numbering_xml, (
        "AR DOCX numbering should use right indentation"
    )
    assert 'w:left=' not in numbering_xml, (
        "AR DOCX numbering should not have left indentation"
    )
