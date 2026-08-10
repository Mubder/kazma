"""DOCX→PDF convert without LibreOffice + PDF sniff honesty."""

from __future__ import annotations

from pathlib import Path

import pytest

from kazma_core.documents.binaries import find_soffice
from kazma_core.documents.config import DocumentConfig
from kazma_core.documents.hostile_corpus import _minimal_pdf
from kazma_core.documents.renderers import RendererReadiness, get_renderer_registry
from kazma_core.documents.service import DocumentService
from kazma_core.documents.sniff import _pdf_has_active_content, sniff_document


def test_resolve_prefers_available_docx_pdf_engine() -> None:
    reg = get_renderer_registry()
    cap = reg.resolve("convert:docx:pdf")
    # With reportlab installed, convert must not stuck on unavailable libreoffice.
    assert cap.available is True
    assert cap.renderer_id in {"libreoffice", "reportlab-office"}
    if find_soffice() is None:
        assert cap.renderer_id == "reportlab-office"


def test_pdf_sniff_allows_openaction_goto(tmp_path: Path) -> None:
    path = tmp_path / "goto.pdf"
    path.write_bytes(
        b"%PDF-1.7\n1 0 obj << /OpenAction << /S /GoTo /D [0 /Fit] >> >>\nendobj\n%%EOF\n"
    )
    assert _pdf_has_active_content(path) is False
    result = sniff_document(path, DocumentConfig(storage_root=tmp_path / "store"))
    assert result.mime_type == "application/pdf"


def test_pdf_sniff_rejects_javascript_openaction(tmp_path: Path) -> None:
    path = tmp_path / "js.pdf"
    path.write_bytes(
        b"%PDF-1.7\n1 0 obj << /OpenAction << /S /JavaScript "
        b"/JS (app.launchURL) >> >>\nendobj\n%%EOF\n"
    )
    assert _pdf_has_active_content(path) is True
    with pytest.raises(Exception) as exc:
        sniff_document(path, DocumentConfig(storage_root=tmp_path / "store"))
    assert "JavaScript" in str(exc.value) or "disabled" in str(exc.value).lower()


def test_pdf_sniff_accepts_minimal_blank_pdf(tmp_path: Path) -> None:
    path = tmp_path / "blank.pdf"
    path.write_bytes(_minimal_pdf(1))
    result = sniff_document(path, DocumentConfig(storage_root=tmp_path / "store"))
    assert result.mime_type == "application/pdf"


@pytest.mark.asyncio
async def test_convert_docx_to_pdf_without_libreoffice(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    pytest.importorskip("reportlab")
    from docx import Document

    root = tmp_path / "store"
    cfg = DocumentConfig(storage_root=root)
    svc = DocumentService(config=cfg)
    source = tmp_path / "note.docx"
    document = Document()
    document.add_heading("Convert Title", level=1)
    document.add_paragraph("Body text for the fallback Office→PDF path.")
    document.save(source)

    result = await svc.convert(
        source,
        "pdf",
        approved_path=source,
        tenant_id="tenant",
        workspace_id="ws",
        actor_id="actor",
    )
    assert result.ok is True, (result.code, result.message)
    assert result.data is not None
    assert result.data.manifest.renderer in {"libreoffice", "reportlab-office"}
    assert result.data.manifest.output_extension == ".pdf"
    assert result.data.storage_path.is_file()
    assert result.data.manifest.output_size > 0
    if result.data.manifest.renderer == "reportlab-office":
        assert any("text extraction" in w.lower() for w in result.data.manifest.warnings)


def test_reportlab_office_capability_ready() -> None:
    reg = get_renderer_registry()
    office = next(c for c in reg.capabilities() if c.renderer_id == "reportlab-office")
    assert office.readiness is RendererReadiness.READY
    assert "convert:docx:pdf" in office.operations


@pytest.mark.asyncio
async def test_delete_document_soft_archives(tmp_path: Path) -> None:
    """Delete tombstones the doc so it no longer appears in the library list."""
    import asyncio
    import io

    from kazma_core.documents.ingestion import DocumentIngestionService

    cfg = DocumentConfig(
        storage_root=tmp_path / "store",
        worker_concurrency=1,
        worker_lease_seconds=5,
        worker_heartbeat_seconds=1,
    )
    svc = DocumentIngestionService(config=cfg)
    await svc.start_workers()
    try:
        result = await asyncio.to_thread(
            svc.ingest_stream,
            io.BytesIO(b"Archive me.\n"),
            filename="archive-me.txt",
            tenant_id="tenant-a",
            workspace_id="ws",
            actor_id="alice",
        )
        loop = asyncio.get_running_loop()
        deadline = loop.time() + 20.0
        while loop.time() < deadline:
            status = await asyncio.to_thread(
                svc.job_status, tenant_id="tenant-a", job_id=result.job_id
            )
            if status and status["state"] == "ready":
                break
            if status and status["state"] in {"rejected", "dead_letter", "cancelled"}:
                raise AssertionError(f"ingest ended in {status['state']}")
            await asyncio.sleep(0.05)
        else:
            raise AssertionError("ingest did not become ready")

        listed = svc.list_documents(tenant_id="tenant-a", actor_id="alice")
        assert any(d["document_id"] == str(result.document_id) for d in listed)

        deleted = svc.delete_document(
            tenant_id="tenant-a",
            actor_id="alice",
            document_id=result.document_id,
            reason="user_requested",
        )
        assert deleted["deleted"] is True

        after = svc.list_documents(tenant_id="tenant-a", actor_id="alice")
        assert not any(d["document_id"] == str(result.document_id) for d in after)
    finally:
        await svc.stop_workers()
        svc.close()
