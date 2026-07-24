"""Document Generator Native Skill — create PDF/DOCX/XLSX/Markdown files.

Each tool lazily imports its heavy dependency and returns a friendly
install-hint string when the library is missing, so the skill always loads.
Output is written to ``kazma-data/documents/``.
"""

from __future__ import annotations

import logging
import re
import time
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

DOC_DIR = Path("kazma-data/documents")


def _slugify(text: str, max_len: int = 50) -> str:
    slug = re.sub(r"[^a-z0-9\s-]", "", text.lower().strip())
    slug = re.sub(r"[\s-]+", "-", slug).strip("-")
    return (slug or "document")[:max_len]


def _filename(title: str, ext: str) -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    ts = int(time.time())
    return DOC_DIR / f"{ts}_{_slugify(title)}.{ext}"


def _normalize_sections(sections: list[Any]) -> list[dict[str, str]]:
    """Coerce a sections list into [{heading, body}] dicts."""
    out: list[dict[str, str]] = []
    for s in sections or []:
        if isinstance(s, dict):
            out.append({
                "heading": str(s.get("heading", "")),
                "body": str(s.get("body", "")),
            })
        elif isinstance(s, (list, tuple)) and len(s) >= 2:
            out.append({"heading": str(s[0]), "body": str(s[1])})
    return out


async def generate_pdf(
    title: str,
    sections: list[dict[str, str]],
) -> str:
    """Generate a PDF from a title and sections (each {heading, body}).

    Requires the ``reportlab`` package (``pip install reportlab``).
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_RIGHT
    except ImportError:
        return "Error: reportlab not installed. Run: pip install reportlab"

    # Arabic text needs reshaping + bidi for proper rendering in reportlab.
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        _has_arabic = True
    except ImportError:
        _has_arabic = False

    def _shape(text: str) -> str:
        """Reshape Arabic text for reportlab (RTL + ligatures)."""
        if not _has_arabic:
            return text
        try:
            return get_display(arabic_reshaper.reshape(text))
        except Exception:
            return text

    # Register fonts. Calibri is the default for the Kazma UI (16px base).
    # For Arabic, Calibri has limited glyphs so we fall back to Amiri/Noto
    # which have full Arabic + ligature support.
    _font_dir = Path(__file__).parent / "fonts"
    _calibri_regular = _font_dir / "calibri.ttf"
    _calibri_bold = _font_dir / "calibrib.ttf"
    _calibri_italic = _font_dir / "calibrii.ttf"

    # Try Calibri first (matches the UI), then fall back to Arabic-capable fonts.
    _arabic_font_name = None
    _font_candidates = [
        _font_dir / "Amiri-Regular.ttf",
        Path("/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf"),
        Path("/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoNaskhArabic-Regular.otf"),
        Path("/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf"),
    ]
    for _fp in _font_candidates:
        if _fp.exists():
            try:
                pdfmetrics.registerFont(TTFont("ArabicFont", str(_fp)))
                _arabic_font_name = "ArabicFont"
                logger.info("[PDF] Registered Arabic font: %s", _fp)
                break
            except Exception:
                continue

    # If no Arabic font found, try downloading Amiri from GitHub.
    if _arabic_font_name is None and _has_arabic:
        try:
            import urllib.request
            _dl_dir = Path("kazma-data/fonts")
            _dl_dir.mkdir(parents=True, exist_ok=True)
            _dl_path = _dl_dir / "Amiri-Regular.ttf"
            if not _dl_path.exists():
                _url = "https://github.com/aliftype/amiri/raw/main/fonts/ttf/Amiri-Regular.ttf"
                logger.info("[PDF] Downloading Amiri Arabic font...")
                urllib.request.urlretrieve(_url, str(_dl_path))
            pdfmetrics.registerFont(TTFont("ArabicFont", str(_dl_path)))
            _arabic_font_name = "ArabicFont"
        except Exception as exc:
            logger.warning("[PDF] Could not obtain Arabic font: %s", exc)

    # Register Calibri for Latin text (if available).
    _calibri_name = None
    if _calibri_regular.exists():
        try:
            from reportlab.pdfbase.pdfmetrics import registerFontFamily
            pdfmetrics.registerFont(TTFont("Calibri", str(_calibri_regular)))
            if _calibri_bold.exists():
                pdfmetrics.registerFont(TTFont("Calibri-Bold", str(_calibri_bold)))
            if _calibri_italic.exists():
                pdfmetrics.registerFont(TTFont("Calibri-Italic", str(_calibri_italic)))
            registerFontFamily(
                "Calibri",
                normal="Calibri",
                bold="Calibri-Bold" if _calibri_bold.exists() else "Calibri",
                italic="Calibri-Italic" if _calibri_italic.exists() else "Calibri",
                boldItalic="Calibri-Bold" if _calibri_bold.exists() else "Calibri",
            )
            _calibri_name = "Calibri"
            logger.info("[PDF] Registered Calibri font family")
        except Exception as exc:
            logger.warning("[PDF] Could not register Calibri: %s", exc)

    secs = _normalize_sections(sections)
    dest = _filename(title, "pdf")
    try:
        doc = SimpleDocTemplate(str(dest), pagesize=LETTER)

        if _arabic_font_name:
            _base_font = _calibri_name or _arabic_font_name
            title_style = ParagraphStyle("Title", fontName=_arabic_font_name, fontSize=18, alignment=TA_RIGHT, spaceAfter=12)
            head_style = ParagraphStyle("SectionHead", fontName=_arabic_font_name, fontSize=14, alignment=TA_RIGHT, spaceBefore=14, spaceAfter=6)
            body_style = ParagraphStyle("BodyText", fontName=_arabic_font_name, fontSize=10, alignment=TA_RIGHT, leading=16)
        elif _calibri_name:
            title_style = ParagraphStyle("Title", fontName="Calibri", fontSize=18, spaceAfter=12)
            head_style = ParagraphStyle("SectionHead", fontName="Calibri-Bold", fontSize=14, spaceBefore=14, spaceAfter=6)
            body_style = ParagraphStyle("BodyText", fontName="Calibri", fontSize=10, leading=16)
        else:
            styles = getSampleStyleSheet()
            title_style = styles["Title"]
            head_style = ParagraphStyle("SectionHead", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
            body_style = styles["BodyText"]

        flow: list[Any] = [Paragraph(_shape(title), title_style), Spacer(1, 12)]
        for s in secs:
            if s["heading"]:
                flow.append(Paragraph(_shape(s["heading"]), head_style))
            if s["body"]:
                flow.append(Paragraph(_shape(s["body"]).replace("\n", "<br/>"), body_style))
            flow.append(Spacer(1, 8))
        doc.build(flow)
    except Exception as exc:  # noqa: BLE001
        return f"Error: PDF generation failed — {exc}"
    return f"PDF generated successfully.\n  Title: {title}\n  Sections: {len(secs)}\n  Saved to: {dest}"


async def generate_docx(
    title: str,
    sections: list[dict[str, str]],
) -> str:
    """Generate a Word .docx from a title and sections (each {heading, body}).

    Requires ``python-docx`` (``pip install python-docx``).
    """
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    secs = _normalize_sections(sections)
    dest = _filename(title, "docx")
    try:
        doc = Document()
        # Set Calibri as the default font (matches the Kazma UI).
        try:
            from docx.shared import Pt
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
        except Exception:
            pass  # Non-fatal — default font still works
        doc.add_heading(title, level=0)
        for s in secs:
            if s["heading"]:
                # Detect heading level from markdown ## prefixes.
                heading = s["heading"].lstrip("#").strip()
                level = s["heading"].count("#") if s["heading"].startswith("#") else 1
                level = min(level, 3)  # docx supports 0-3 practically
                doc.add_heading(heading, level=level)
            if s["body"]:
                # Split body into paragraphs and add each separately.
                for para in s["body"].split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
        doc.save(str(dest))
    except Exception as exc:  # noqa: BLE001
        return f"Error: DOCX generation failed — {exc}"
    return f"DOCX generated successfully.\n  Title: {title}\n  Sections: {len(secs)}\n  Saved to: {dest}"


async def generate_xlsx(
    sheets: list[dict[str, Any]],
    filename: str = "workbook",
) -> str:
    """Generate an Excel .xlsx from sheets (each {name, rows}).

    ``rows`` is a list of row-lists; the first row is treated as the header.
    Requires ``openpyxl`` (``pip install openpyxl``).
    """
    try:
        from openpyxl import Workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    dest = _filename(filename, "xlsx")
    try:
        wb = Workbook()
        # Remove the default sheet; we add named ones below.
        wb.remove(wb.active)
        for sh in sheets or []:
            name = str(sh.get("name", "Sheet"))[:31]
            rows = sh.get("rows") or []
            ws = wb.create_sheet(title=name)
            for row in rows:
                ws.append([str(c) if c is not None else "" for c in row])
        if not wb.sheetnames:
            wb.create_sheet("Sheet")
        wb.save(str(dest))
    except Exception as exc:  # noqa: BLE001
        return f"Error: XLSX generation failed — {exc}"
    return (
        f"XLSX generated successfully.\n"
        f"  Sheets: {len(sheets or [])}\n"
        f"  Saved to: {dest}"
    )


async def generate_markdown_doc(
    title: str,
    sections: list[dict[str, str]],
) -> str:
    """Generate a Markdown (.md) document from a title and sections.

    No external dependency required.
    """
    secs = _normalize_sections(sections)
    dest = _filename(title, "md")
    try:
        lines: list[str] = [f"# {title}", ""]
        for s in secs:
            if s["heading"]:
                lines.append(f"## {s['heading']}")
                lines.append("")
            if s["body"]:
                lines.append(s["body"])
                lines.append("")
        dest.write_text("\n".join(lines), encoding="utf-8")
    except OSError as exc:
        return f"Error: could not write markdown file — {exc}"
    return f"Markdown generated successfully.\n  Title: {title}\n  Sections: {len(secs)}\n  Saved to: {dest}"
